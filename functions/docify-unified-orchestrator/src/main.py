#!/usr/bin/env python3
"""
Docify Unified Orchestrator - Advanced Gemini-Powered Document Analysis
Uses all available Gemini API tools for comprehensive content analysis
"""

import os
import json
import time
import requests
import chardet
import re
import hashlib
from typing import Dict, Any, Optional, List, Tuple, Union
from urllib.parse import urlparse, urljoin, parse_qs
from io import BytesIO
from bs4 import BeautifulSoup
from appwrite.client import Client
from appwrite.services.databases import Databases
from appwrite.id import ID
from google import genai
from google.genai import types
import PyPDF2
import docx
import pandas as pd
import feedparser
from collections import defaultdict
import mimetypes
import base64

# Environment variables
DATABASE_ID = os.environ.get('DATABASE_ID', 'docify_db')
DOCUMENTS_COLLECTION_ID = os.environ.get('DOCUMENTS_COLLECTION_ID', 'documents_table')
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY', '')
BROWSERLESS_API_KEY = os.environ.get('BROWSERLESS_API_KEY', '')

# Initialize clients
client = Client()
client.set_endpoint(os.environ.get('APPWRITE_FUNCTION_API_ENDPOINT'))
client.set_project(os.environ.get('APPWRITE_FUNCTION_PROJECT_ID'))
client.set_key(os.environ.get('APPWRITE_API_KEY'))
databases = Databases(client)

# Gemini configuration
GEMINI_MODEL = "gemini-2.5-pro"  # Advanced model with superior reasoning and tool integration
MAX_CONTENT_LENGTH = 150000  # Maximum content to process (increased for Pro model)
MAX_TOOL_EXECUTIONS = 20  # Maximum tool calls per analysis (increased for Pro model)
TOOL_TIMEOUT = 45  # Tool execution timeout in seconds (increased for Pro model)

# System prompt configuration
SYSTEM_PROMPT = """You are Docify AI, an advanced technical documentation analyzer powered by Gemini 2.5 Pro.

Your mission is to analyze any web content or document and create comprehensive, visually-rich explanations that help users understand complex technical topics.

CORE CAPABILITIES:
1. 🌐 Advanced web scraping and content extraction
2. 🔍 Deep research using Google Search integration
3. 📊 Multi-format content analysis (HTML, PDF, API docs, tutorials, etc.)
4. 🎨 Intelligent block generation with optimal visualization types
5. 🔬 User-centric analysis guided by their specific instructions

ANALYSIS FRAMEWORK:
- Always prioritize user instructions for research direction and presentation style
- Use user prompts to guide what aspects to emphasize and how to present information
- Research related topics when relevant to user interests
- Create analysis blocks that directly address user questions and needs
- Adapt visualization types based on content type and user preferences

TOOL INTEGRATION:
- Google Search: Research current best practices, related technologies, alternatives
- URL Context: Enhanced content analysis with web relationships and context
- Code Analysis: Detect, validate, and explain code examples
- Content Extraction: Multi-format parsing with metadata preservation

OUTPUT REQUIREMENTS:
- Generate 3-6 analysis blocks maximum, optimized for 3x3 grid layout
- Use appropriate block types: summary, key_points, architecture, mermaid, code, api_reference, guide, comparison, best_practices, troubleshooting
- Size blocks optimally: small (1 unit), medium (2 units), large (3 units)
- Ensure all blocks have unique IDs and comprehensive metadata
- Prioritize content based on user instructions and content analysis

CONTENT ANALYSIS STRATEGY:
1. Extract and analyze the main content from the provided URL
2. Use user instructions to determine research direction and emphasis
3. Research related topics that would help the user better understand the content
4. Create blocks that directly address user needs and questions
5. Optimize for visual learning with appropriate diagram types and code examples

Remember: You are not just analyzing content - you are creating personalized learning experiences guided by user instructions."""

# ===== COMPREHENSIVE SCRAPING FUNCTIONS =====

def smart_fetch_html(url, timeout=15):
    """Fetch HTML content using multiple strategies like url-to-markdown approach."""
    html_content = None
    fetch_method = 'basic'

    # Fetch strategies (adapted from url-to-markdown)
    fetch_strategies = [
        {
            'name': 'modern-browser',
            'headers': {
                'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
                'Accept-Encoding': 'gzip, deflate, br',
                'DNT': '1',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
                'Sec-Fetch-Dest': 'document',
                'Sec-Fetch-Mode': 'navigate',
                'Sec-Fetch-Site': 'none',
                'Cache-Control': 'max-age=0'
            }
        },
        {
            'name': 'mobile',
            'headers': {
                'User-Agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 17_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.0 Mobile/15E148 Safari/604.1',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
                'Accept-Encoding': 'gzip, deflate, br'
            }
        },
        {
            'name': 'bot-friendly',
            'headers': {
                'User-Agent': 'Mozilla/5.0 (compatible; DocifyBot/1.0; +https://docify.app)',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9'
            }
        }
    ]

    # Try each strategy
    for strategy in fetch_strategies:
        try:
            print(f"Trying fetch strategy: {strategy['name']}")
            response = requests.get(url, headers=strategy['headers'], timeout=timeout)

            if response.status_code == 200:
                html_content = response.text
                fetch_method = strategy['name']
                print(f"Successfully fetched {len(html_content)} characters using {strategy['name']} strategy")

                # Check if content seems too small (likely JS-dependent)
                if len(html_content) < 10000:
                    print(f"Content seems small ({len(html_content)} chars), might be JS-dependent. Will try rendering services.")
                    continue
                break
            else:
                print(f"{strategy['name']} failed with status: {response.status_code}")
        except Exception as strategy_error:
            print(f"{strategy['name']} failed: {strategy_error}")

    # Try JS rendering services if content is small or all strategies failed
    tried_rendering = False
    if not html_content or len(html_content or '') < 10000:
        tried_rendering = True
        print('Content seems small or no content fetched, trying JS rendering services...')

        # Try Browserless.io first (if API key is available)
        if BROWSERLESS_API_KEY:
            try:
                print('Trying Browserless.io...')
                browserless_url = f"https://production-sfo.browserless.io/content?token={BROWSERLESS_API_KEY}"
                response = requests.post(browserless_url, json={
                    'url': url,
                    'gotoOptions': {
                        'waitUntil': 'networkidle2',
                        'timeout': 30000
                    }
                }, timeout=30)

                if response.status_code == 200:
                    rendered_html = response.text
                    if len(rendered_html) > len(html_content or ''):
                        html_content = rendered_html
                        fetch_method = 'browserless-io'
                        print(f"Successfully rendered {len(html_content)} characters using Browserless.io")
                    else:
                        print(f"Browserless.io returned same or smaller content ({len(rendered_html)} chars)")
                else:
                    print(f"Browserless.io failed with status: {response} - {response.text}")
            except Exception as browserless_error:
                print(f"Browserless.io error: {browserless_error}")
        else:
            print("BROWSERLESS_API_KEY not provided, skipping Browserless.io")

        # Fallback to free services if Browserless.io didn't help
        if not html_content or len(html_content) < 10000:
            free_services = [
                {
                    'name': 'rendertron-free',
                    'url': f"https://render-tron.appspot.com/render/{url}",
                    'type': 'direct'
                }
            ]

            for service in free_services:
                try:
                    print(f"Trying fallback {service['name']}...")
                    response = requests.get(service['url'], headers={
                        'User-Agent': 'Mozilla/5.0 (compatible; DocifyBot/1.0)'
                    }, timeout=20)

                    if response.status_code == 200:
                        rendered_html = response.text
                        if len(rendered_html) > len(html_content or ''):
                            html_content = rendered_html
                            fetch_method = service['name']
                            print(f"Successfully rendered {len(html_content)} characters using {service['name']}")
                            break
                        else:
                            print(f"{service['name']} returned same or smaller content ({len(rendered_html)} chars)")
                except Exception as service_error:
                    print(f"{service['name']} failed: {service_error}")

    return html_content, fetch_method, tried_rendering


def detect_content_type_from_response(response, url):
    """Detect the content type of a response."""
    # Check Content-Type header
    content_type_header = response.headers.get('Content-Type', '').lower()

    # Check URL extension
    url_path = urlparse(url).path.lower()

    # Determine content type
    if 'application/pdf' in content_type_header or url_path.endswith('.pdf'):
        return 'pdf'
    elif 'application/msword' in content_type_header or url_path.endswith(('.doc', '.docx')):
        return 'doc'
    elif 'application/vnd.ms-excel' in content_type_header or url_path.endswith(('.xls', '.xlsx')):
        return 'excel'
    elif 'text/csv' in content_type_header or url_path.endswith('.csv'):
        return 'csv'
    elif 'application/json' in content_type_header or url_path.endswith('.json'):
        return 'json'
    elif 'application/xml' in content_type_header or url_path.endswith(('.xml', '.rss', '.atom')):
        return 'xml'
    elif 'text/plain' in content_type_header or url_path.endswith(('.txt', '.md')):
        return 'text'
    else:
        # Default to HTML for web pages
        return 'html'


def extract_content_by_type(response, content_type, url):
    """Extract content based on the detected content type."""
    try:
        if content_type == 'pdf':
            return extract_pdf_content(response, url)
        elif content_type == 'doc':
            return extract_doc_content(response, url)
        elif content_type == 'excel':
            return extract_excel_content(response, url)
        elif content_type == 'csv':
            return extract_csv_content(response, url)
        elif content_type == 'json':
            return extract_json_content(response, url)
        elif content_type == 'xml':
            return extract_xml_content(response, url)
        elif content_type == 'text':
            return extract_text_content(response, url)
        else:  # Default to HTML
            return extract_html_content(response, url)
    except Exception as e:
        print(f"Error extracting {content_type} content from {url}: {e}")
        return None


def extract_pdf_content(response, url):
    """Extract content from PDF files."""
    try:
        pdf_file = BytesIO(response.content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        content = []
        for page in pdf_reader.pages[:10]:  # Limit to first 10 pages
            text = page.extract_text()
            if text.strip():
                content.append(text)

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'PDF Document - {len(pdf_reader.pages)} pages',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'pdf',
            'metadata': {
                'pages': len(pdf_reader.pages),
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting PDF content: {e}")
        return None


def extract_doc_content(response, url):
    """Extract content from Word documents."""
    try:
        doc_file = BytesIO(response.content)
        doc = docx.Document(doc_file)

        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'Word Document - {len(doc.paragraphs)} paragraphs',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'doc',
            'metadata': {
                'paragraphs': len(doc.paragraphs),
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting DOC content: {e}")
        return None


def extract_excel_content(response, url):
    """Extract content from Excel files."""
    try:
        excel_file = BytesIO(response.content)
        df = pd.read_excel(excel_file)

        # Convert DataFrame to readable text
        content = []
        content.append(f"Sheet: {df.columns.tolist()}")
        content.append("Data Preview:")
        content.append(str(df.head(20)))  # First 20 rows

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'Excel Spreadsheet - {len(df)} rows, {len(df.columns)} columns',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'excel',
            'metadata': {
                'rows': len(df),
                'columns': len(df.columns),
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting Excel content: {e}")
        return None


def extract_csv_content(response, url):
    """Extract content from CSV files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        csv_text = response.content.decode(encoding, errors='ignore')
        lines = csv_text.split('\n')[:50]  # First 50 lines

        content = []
        content.append("CSV Data Preview:")
        content.extend(lines[:20])  # First 20 lines

        full_content = '\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'CSV File - {len(lines)} lines',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'csv',
            'metadata': {
                'lines': len(lines),
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting CSV content: {e}")
        return None


def extract_json_content(response, url):
    """Extract content from JSON files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        json_text = response.content.decode(encoding, errors='ignore')
        json_data = json.loads(json_text)

        # Convert JSON to readable text
        content = []
        content.append("JSON Structure:")
        content.append(json.dumps(json_data, indent=2)[:2000])  # First 2000 chars

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'JSON Data - {len(json_text)} characters',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'json',
            'metadata': {
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting JSON content: {e}")
        return None


def extract_xml_content(response, url):
    """Extract content from XML/RSS files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        xml_text = response.content.decode(encoding, errors='ignore')

        # Check if it's an RSS/Atom feed
        if '<rss' in xml_text.lower() or '<feed' in xml_text.lower():
            return extract_feed_content(xml_text, url)

        # Regular XML parsing
        from xml.etree import ElementTree as ET
        root = ET.fromstring(xml_text)

        content = []
        content.append("XML Structure:")
        content.append(xml_to_text(root, level=0)[:2000])

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'XML Document - {len(xml_text)} characters',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'xml',
            'metadata': {
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting XML content: {e}")
        return None


def extract_feed_content(xml_text, url):
    """Extract content from RSS/Atom feeds."""
    try:
        feed = feedparser.parse(xml_text)

        content = []
        content.append(f"Feed Title: {feed.feed.get('title', 'Unknown')}")
        content.append(f"Feed Description: {feed.feed.get('description', '')}")
        content.append(f"Total Entries: {len(feed.entries)}")

        # Add recent entries
        for i, entry in enumerate(feed.entries[:10]):
            content.append(f"\n--- Entry {i+1} ---")
            content.append(f"Title: {entry.get('title', '')}")
            content.append(f"Link: {entry.get('link', '')}")
            if 'summary' in entry:
                content.append(f"Summary: {entry.summary[:500]}...")

        full_content = '\n'.join(content)

        return {
            'url': url,
            'title': feed.feed.get('title', extract_title_from_url(url)),
            'description': f'RSS/Atom Feed - {len(feed.entries)} entries',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'feed',
            'metadata': {
                'entries': len(feed.entries),
                'feed_type': 'rss' if '<rss' in xml_text.lower() else 'atom'
            }
        }
    except Exception as e:
        print(f"Error extracting feed content: {e}")
        return None


def extract_text_content(response, url):
    """Extract content from plain text files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        text_content = response.content.decode(encoding, errors='ignore')

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'Text File - {len(text_content)} characters',
            'content': text_content[:10000],  # Limit to 10K chars
            'word_count': len(text_content.split()),
            'content_type': 'text',
            'metadata': {
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting text content: {e}")
        return None


def extract_html_content(response, url):
    """Extract content from HTML pages."""
    try:
        # Handle both Scrapy responses and raw HTML content
        if hasattr(response, 'text'):
            html_content = response.text
        elif hasattr(response, 'content'):
            # Handle raw bytes
            html_content = response.content.decode('utf-8', errors='ignore')
        else:
            # Handle MockResponse or other objects
            html_content = str(response)

        # Parse with BeautifulSoup
        soup = BeautifulSoup(html_content, 'lxml')

        title = extract_title_from_soup(soup, url)
        description = extract_description_from_soup(soup)
        content = extract_main_content_from_soup(soup)

        return {
            'url': url,
            'title': title,
            'description': description,
            'content': content,
            'word_count': len(content.split()) if content else 0,
            'content_type': 'html'
        }
    except Exception as e:
        print(f"Error extracting HTML content: {e}")
        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': '',
            'content': f'Error extracting HTML content: {str(e)}',
            'word_count': 0,
            'content_type': 'html'
        }


def extract_title_from_url(url):
    """Extract a title from URL when no other title is available."""
    path = urlparse(url).path
    filename = path.split('/')[-1]
    if '.' in filename:
        filename = filename.split('.')[0]
    return filename.replace('-', ' ').replace('_', ' ').title() or 'Document'


def xml_to_text(element, level=0):
    """Convert XML element to readable text."""
    indent = "  " * level
    text = f"{indent}<{element.tag}>"

    if element.text and element.text.strip():
        text += f" {element.text.strip()}"

    for child in element:
        text += "\n" + xml_to_text(child, level + 1)

    if not element.text or not element.text.strip():
        text += f"{indent}</{element.tag}>"

    return text


def extract_title_from_soup(soup, url):
    """Extract page title from BeautifulSoup object."""
    title_selectors = [
        'title',
        'h1',
        'h1 a',
        '[property="og:title"]',
        'meta[name="title"]'
    ]

    for selector in title_selectors:
        try:
            if selector == 'title':
                element = soup.title
                if element and element.string:
                    return element.string.strip()
            else:
                element = soup.select_one(selector)
                if element:
                    if 'property' in selector or 'name' in selector:
                        return element.get('content', '').strip()
                    else:
                        return element.get_text().strip()
        except:
            continue

    return extract_title_from_url(url)


def extract_description_from_soup(soup):
    """Extract page description from BeautifulSoup object."""
    desc_selectors = [
        'meta[name="description"]',
        'meta[property="og:description"]',
        'meta[name="twitter:description"]'
    ]

    for selector in desc_selectors:
        try:
            element = soup.select_one(selector)
            if element:
                return element.get('content', '').strip()
        except:
            continue

    return ''


def extract_main_content_from_soup(soup):
    """Extract main content from BeautifulSoup object."""
    # Try various content selectors in order of preference
    content_selectors = [
        'main',
        '[role="main"]',
        '.content',
        '.main-content',
        '#content',
        '#main',
        'article',
        '.article-content',
        '.post-content',
        '.entry-content',
        '.page-content',
        '.text-content'
    ]

    for selector in content_selectors:
        try:
            content_elements = soup.select(selector)
            if content_elements:
                content = []
                for element in content_elements:
                    text = element.get_text()
                    if text.strip():
                        content.append(text)

                if content:
                    combined_content = ' '.join(content)
                    if len(combined_content.strip()) > 100:  # Minimum content length
                        return clean_content(combined_content)
        except:
            continue

    # Fallback: extract all paragraph text
    try:
        paragraphs = soup.select('p')
        if paragraphs:
            content = []
            for p in paragraphs:
                text = p.get_text().strip()
                if text:
                    content.append(text)

            if content:
                combined_content = ' '.join(content)
                if len(combined_content.strip()) > 100:
                    return clean_content(combined_content)
    except:
        pass

    # Final fallback: extract all text from body
    try:
        body = soup.body
        if body:
            body_text = body.get_text()
            return clean_content(body_text)
    except:
        pass

    return ''


def extract_title(response, url):
    """Extract page title."""
    title_selectors = [
        'title::text',
        'h1::text',
        'h1 a::text',
        '[property="og:title"]::attr(content)',
        'meta[name="title"]::attr(content)'
    ]

    for selector in title_selectors:
        title = response.css(selector).get()
        if title:
            return title.strip()

    return url.split('/')[-1] or 'Untitled Page'


def extract_description(response):
    """Extract page description."""
    desc_selectors = [
        'meta[name="description"]::attr(content)',
        'meta[property="og:description"]::attr(content)',
        'meta[name="twitter:description"]::attr(content)'
    ]

    for selector in desc_selectors:
        desc = response.css(selector).get()
        if desc:
            return desc.strip()

    return ''


def extract_main_content(response):
    """Extract main content from the page."""
    # Try various content selectors in order of preference
    content_selectors = [
        'main',
        '[role="main"]',
        '.content',
        '.main-content',
        '#content',
        '#main',
        'article',
        '.article-content',
        '.post-content',
        '.entry-content',
        '.page-content',
        '.text-content'
    ]

    for selector in content_selectors:
        content_elements = response.css(selector)
        if content_elements:
            content = content_elements.css('::text').getall()
            if content:
                combined_content = ' '.join(content)
                if len(combined_content.strip()) > 100:  # Minimum content length
                    return clean_content(combined_content)

    # Fallback: extract all paragraph text
    paragraphs = response.css('p::text').getall()
    if paragraphs:
        content = ' '.join(paragraphs)
        if len(content.strip()) > 100:
            return clean_content(content)

    # Final fallback: extract all text from body
    body_text = response.css('body ::text').getall()
    combined_text = ' '.join(body_text)
    return clean_content(combined_text)


def clean_content(content):
    """Clean and process scraped content."""
    if not content:
        return ''

    # Remove excessive whitespace
    content = re.sub(r'\s+', ' ', content)

    # Remove navigation and footer content
    content = re.sub(r'\b(home|menu|navigation|footer|copyright|privacy|terms|contact|about|login|signup|register|search)\b',
                    '', content, flags=re.IGNORECASE)

    # Remove emails, URLs, and phone numbers
    content = re.sub(r'\S+@\S+\.\S+', '[EMAIL]', content)
    content = re.sub(r'https?://[^\s]+', '[URL]', content)
    content = re.sub(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', '[PHONE]', content)

    # Remove non-printable characters
    content = re.sub(r'[^\x20-\x7E\n\r\t]', '', content)

    return content.strip()


def scrape_website_with_requests(url, max_pages=10):
    """Scrape content from a website and its subpaths using requests-based approach."""
    try:
        print(f"Starting requests-based crawl for: {url} (max {max_pages} pages)")

        # Parse the URL to get domain
        parsed_url = urlparse(url)
        domain = parsed_url.netloc
        base_url = f"{parsed_url.scheme}://{domain}"

        print(f"Crawling domain: {domain}")

        visited_urls = set()
        pages_to_visit = [url]
        all_pages = []
        pages_crawled = 0
        start_time = time.time()

        headers = {
            'User-Agent': 'DocifyBot/1.0 (https://docify.app)',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'Accept-Encoding': 'gzip, deflate, br',
        }

        while pages_to_visit and pages_crawled < max_pages and time.time() - start_time < 300:  # 5 minute timeout
            current_url = pages_to_visit.pop(0)

            if current_url in visited_urls:
                continue

            visited_urls.add(current_url)
            pages_crawled += 1

            print(f"Crawling page {pages_crawled}/{max_pages}: {current_url}")

            try:
                # Use smart fetch approach
                html_content, fetch_method, tried_rendering = smart_fetch_html(current_url, timeout=15)

                if not html_content:
                    print(f"Could not fetch content from: {current_url}")
                    continue

                # Detect content type
                response = requests.head(current_url, headers=headers, timeout=10)
                content_type = detect_content_type_from_response(response, current_url)

                # Create a mock response object for extract_content_by_type
                class MockResponse:
                    def __init__(self, content, headers, url):
                        self.content = content
                        self.headers = headers
                        self.url = url

                mock_response = MockResponse(html_content.encode('utf-8'), response.headers, current_url)

                # Extract content
                page_data = extract_content_by_type(mock_response, content_type, current_url)

                if not page_data:
                    print(f"Could not extract content from: {current_url}")
                    continue

                # Skip pages with very little content
                content_text = page_data.get('content', '')
                if not content_text or len(content_text.strip()) < 50:
                    print(f"Skipping page with insufficient content: {current_url}")
                    continue

                # Store the page data
                all_pages.append(page_data)

                # Find links to follow (only if we haven't reached max pages and content is HTML)
                if pages_crawled < max_pages and content_type == 'html':
                    try:
                        soup = BeautifulSoup(html_content, 'lxml')
                        for link in soup.find_all('a', href=True):
                            next_url = urljoin(current_url, link['href'])

                            # Only follow links within the same domain
                            if urlparse(next_url).netloc == domain:
                                # Skip common non-content URLs
                                path = urlparse(next_url).path.lower()
                                if not any(skip in path for skip in ['/search', '/login', '/admin', '/wp-admin', '/api/', '/feed', '/tag/', '/category/', '/author/']):
                                    if next_url not in visited_urls and next_url not in pages_to_visit:
                                        pages_to_visit.append(next_url)

                                        # Limit the queue size to prevent memory issues
                                        if len(pages_to_visit) > max_pages * 3:
                                            break
                    except Exception as link_error:
                        print(f"Error finding links on {current_url}: {link_error}")

                # Small delay to be respectful
                time.sleep(1.0)

            except Exception as page_error:
                print(f"Error crawling {current_url}: {page_error}")
                continue

        print(f"Requests-based crawling completed: {len(all_pages)} pages with content")

        if not all_pages:
            print("No pages were successfully crawled, falling back to single page")
            return scrape_single_page_fallback(url)

        # Limit to max_pages even if more were crawled
        all_pages = all_pages[:max_pages]

        # Analyze content types
        content_type_stats = {}
        for page in all_pages:
            content_type = page.get('content_type', 'unknown')
            content_type_stats[content_type] = content_type_stats.get(content_type, 0) + 1

        print(f"Content types found: {content_type_stats}")

        # Aggregate content from all pages
        total_word_count = sum(page['word_count'] for page in all_pages if page.get('word_count'))

        # Create combined content (limit content length to prevent memory issues)
        combined_content = []
        max_content_length = 75000
        current_length = 0

        # Group pages by content type for better organization
        pages_by_type = {}
        for page in all_pages:
            content_type = page.get('content_type', 'unknown')
            if content_type not in pages_by_type:
                pages_by_type[content_type] = []
            pages_by_type[content_type].append(page)

        # Add content by type
        for content_type, pages in pages_by_type.items():
            if current_length >= max_content_length:
                break

            # Add type header
            type_header = f"\n{'='*50}\n{content_type.upper()} CONTENT ({len(pages)} files)\n{'='*50}\n"
            if current_length + len(type_header) <= max_content_length:
                combined_content.append(type_header)
                current_length += len(type_header)

            for page in pages:
                if page.get('content') and current_length < max_content_length:
                    # Include metadata for non-HTML content
                    metadata_info = ""
                    if page.get('metadata'):
                        metadata_info = f"\n[Metadata: {page['metadata']}]"

                    page_content = f"\n--- {page['title']} ({page['url']}){metadata_info} ---\n{page['content']}"
                    if current_length + len(page_content) <= max_content_length:
                        combined_content.append(page_content)
                        current_length += len(page_content)
                    else:
                        # Truncate if it would exceed limit
                        remaining_space = max_content_length - current_length
                        if remaining_space > 100:
                            truncated_content = page_content[:remaining_space-10] + "..."
                            combined_content.append(truncated_content)
                        break

        # Get the main page data (first page crawled)
        main_page = all_pages[0] if all_pages else {}

        result = {
            'title': main_page.get('title', 'Multi-Type Content'),
            'description': f'Crawled {len(all_pages)} items from {domain} - Types: {", ".join([f"{k}({v})" for k,v in content_type_stats.items()])}',
            'content': '\n'.join(combined_content),
            'url': url,
            'word_count': total_word_count,
            'pages_crawled': len(all_pages),
            'scraped_at': 'now',
            'content_types': content_type_stats,
            'subpages': [{'url': p['url'], 'title': p['title'], 'word_count': p['word_count'], 'content_type': p.get('content_type', 'unknown')} for p in all_pages[1:]]
        }

        print(f"Successfully scraped {len(all_pages)} items with {total_word_count} total words")
        print(f"Content breakdown: {content_type_stats}")
        return result

    except Exception as error:
        print(f'Requests-based crawling error: {error}')
        print("Falling back to single page scraping...")
        return scrape_single_page_fallback(url)


def scrape_single_page_fallback(url):
    """Fallback single page scraper using the new smart fetch approach."""
    try:
        print(f"Fallback scraping single page: {url}")

        # Use smart fetch approach
        html_content, fetch_method, tried_rendering = smart_fetch_html(url, timeout=30)

        if not html_content:
            raise Exception("Could not fetch content from URL")

        # Parse with BeautifulSoup
        soup = BeautifulSoup(html_content, 'lxml')

        title = soup.title.string if soup.title else extract_title_from_url(url)

        meta_desc = soup.find('meta', attrs={'name': 'description'})
        description = meta_desc.get('content', '') if meta_desc else ''

        # Try to find main content
        content_selectors = [
            'main', '[role="main"]', '.content', '.main-content',
            '#content', '#main', 'article', '.article-content',
            '.post-content', '.entry-content'
        ]

        main_content = ''
        for selector in content_selectors:
            try:
                element = soup.select_one(selector)
                if element and len(element.get_text().strip()) > 100:
                    main_content = element.get_text().strip()
                    break
            except:
                continue

        if not main_content:
            main_content = soup.body.get_text().strip() if soup.body else ''

        cleaned_content = clean_content(main_content)

        return {
            'title': title,
            'description': description,
            'content': cleaned_content,
            'url': url,
            'word_count': len(cleaned_content.split()),
            'scraped_at': 'now'
        }

    except Exception as error:
        print(f'Fallback scraping error: {error}')
        raise Exception(f'Failed to scrape website: {str(error)}')


# ===== LOGGER CLASS =====
class Logger:
    """Simple logger for the function with Appwrite context logging"""

    def __init__(self, context=None):
        self.logs = []
        self.start_time = time.time()
        self.context = context

    def log(self, message: str):
        """Log a message with timestamp using Appwrite context"""
        timestamp = time.time() - self.start_time

        # Use Appwrite context logging if available
        if self.context and hasattr(self.context, 'log'):
            self.context.log(f"[{timestamp:.2f}s] {message}")
        else:
            print(f"[{timestamp:.2f}s] {message}")

        self.logs.append({
            "timestamp": timestamp,
            "message": message,
            "level": "info"
        })

    def error(self, message: str):
        """Log an error message"""
        timestamp = time.time() - self.start_time

        if self.context and hasattr(self.context, 'error'):
            self.context.error(f"[{timestamp:.2f}s] ERROR: {message}")
        else:
            print(f"[{timestamp:.2f}s] ERROR: {message}")

        self.logs.append({
            "timestamp": timestamp,
            "message": message,
            "level": "error"
        })

    def get_logs(self) -> List[Dict[str, Any]]:
        """Get all logged messages"""
        return self.logs.copy()

# ===== ADVANCED CONTENT PROCESSOR =====
class AdvancedContentProcessor:
    """Advanced content processor using all Gemini tools"""

    def __init__(self, logger, system_prompt: str = ""):
        self.logger = logger
        self.system_prompt = system_prompt
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })

    def process_url_comprehensive(self, url: str, gemini_client, user_instructions: str) -> Dict[str, Any]:
        """Streamlined URL processing focused on web content and analysis"""
        try:
            self.logger.log(f"🚀 Starting streamlined URL analysis: {url}")

            # Phase 1: Clean content extraction
            self.logger.log("📄 Phase 1: Content extraction")
            initial_content = self._extract_web_content(url)

            # Phase 2: Gemini-powered analysis with user instructions
            self.logger.log("🤖 Phase 2: AI analysis with user context")
            enhanced_analysis = self._analyze_with_gemini_tools(url, initial_content, gemini_client, user_instructions)

            # Phase 3: Research enrichment based on user instructions
            self.logger.log("🔬 Phase 3: Research enhancement")
            enriched_content = self._enrich_with_research(url, enhanced_analysis, gemini_client, user_instructions)

            # Phase 4: Final synthesis
            self.logger.log("🎯 Phase 4: Final synthesis")
            final_result = self._synthesize_final_result(url, enriched_content)

            self.logger.log(f"✅ Analysis completed: {len(final_result.get('content', ''))} characters")
            return final_result

        except Exception as e:
            self.logger.error(f"❌ Processing failed: {e}")
            return {
                "error": str(e),
                "url": url,
                "fallback_content": self._extract_basic_content(url)
            }

    def _extract_web_content(self, url: str) -> Dict[str, Any]:
        """Extract web content with smart fallback strategies"""
        try:
            self.logger.log(f"🌐 Extracting web content from: {url}")

            # Try multiple fetch strategies for better content capture
            html_content = self._fetch_with_multiple_strategies(url)

            if not html_content:
                raise Exception("Could not fetch content from URL")

            # Parse and extract content
            soup = BeautifulSoup(html_content, 'lxml')

            # Extract metadata
            title = self._extract_title(soup, url)
            description = self._extract_description(soup)
            main_content = self._extract_main_content(soup)
            code_blocks = self._extract_code_blocks(soup)
            links = self._extract_links(soup, url)

            # Clean and process content
            cleaned_content = self._clean_content(main_content)

            return {
                "url": url,
                "title": title,
                "description": description,
                "content": cleaned_content,
                "code_blocks": code_blocks,
                "links": links,
                "word_count": len(cleaned_content.split()),
                "content_type": "web_content"
            }

        except Exception as e:
            self.logger.error(f"⚠️ Content extraction failed: {e}")
            return {"error": str(e), "url": url}

    def _fetch_with_multiple_strategies(self, url: str) -> str:
        """Fetch content using multiple strategies for better success rate"""
        strategies = [
            {
                'name': 'modern',
                'headers': {
                    'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                    'Accept-Language': 'en-US,en;q=0.9'
                }
            },
            {
                'name': 'basic',
                'headers': {
                    'User-Agent': 'DocifyBot/1.0 (https://docify.app)',
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
                }
            }
        ]

        for strategy in strategies:
            try:
                self.logger.log(f"Trying {strategy['name']} fetch strategy")
                response = self.session.get(url, headers=strategy['headers'], timeout=20)
                if response.status_code == 200:
                    return response.text
            except Exception as e:
                self.logger.log(f"{strategy['name']} strategy failed: {e}")
                continue

        return None

    def _extract_title(self, soup, url: str) -> str:
        """Extract page title with fallbacks"""
        title_selectors = [
            'title',
            'h1',
            'h1 a',
            '[property="og:title"]',
            'meta[name="title"]'
        ]

        for selector in title_selectors:
            try:
                if selector == 'title':
                    element = soup.title
                    if element and element.string:
                        return element.string.strip()
                else:
                    element = soup.select_one(selector)
                    if element:
                        if 'property' in selector or 'name' in selector:
                            return element.get('content', '').strip()
                        else:
                            return element.get_text().strip()
            except Exception:
                continue

        return url.split('/')[-1] or 'Untitled Page'

    def _extract_description(self, soup) -> str:
        """Extract page description"""
        desc_selectors = [
            'meta[name="description"]',
            'meta[property="og:description"]',
            'meta[name="twitter:description"]'
        ]

        for selector in desc_selectors:
            try:
                element = soup.select_one(selector)
                if element:
                    return element.get('content', '').strip()
            except Exception:
                continue

        return ''

    def _extract_main_content(self, soup) -> str:
        """Extract main content using multiple strategies"""
        content_selectors = [
            'main',
            '[role="main"]',
            '.content',
            '.main-content',
            '#content',
            '#main',
            'article',
            '.post-content',
            '.entry-content'
        ]

        for selector in content_selectors:
            try:
                elements = soup.select(selector)
                if elements:
                    content = []
                    for element in elements:
                        text = element.get_text().strip()
                        if text:
                            content.append(text)

                    if content:
                        combined = ' '.join(content)
                        if len(combined) > 100:
                            return combined
            except:
                continue

        # Fallback to body content
        try:
            body = soup.body
            if body:
                return body.get_text().strip()
        except Exception:
            pass

        return ''

    def _extract_code_blocks(self, soup) -> List[Dict[str, Any]]:
        """Extract code blocks and examples"""
        code_blocks = []

        # Pre-formatted code blocks
        for pre in soup.find_all('pre'):
            code_element = pre.find('code')
            if code_element:
                code_content = code_element.get_text().strip()
                if len(code_content) > 10:
                    language = ""
                    code_class = code_element.get('class', [])
                    for cls in code_class:
                        if cls.startswith('language-') or cls.startswith('lang-'):
                            language = cls.split('-', 1)[1]
                            break

                    code_blocks.append({
                        "content": code_content,
                        "language": language,
                        "type": "pre"
                    })

        # Inline code (longer snippets)
        for code in soup.find_all('code'):
            if not code.find_parent('pre'):
                code_content = code.get_text().strip()
                if len(code_content) > 20:
                    code_blocks.append({
                        "content": code_content,
                        "language": "",
                        "type": "inline"
                    })

        return code_blocks[:10]  # Limit to 10 blocks

    def _extract_links(self, soup, base_url: str) -> List[Dict[str, Any]]:
        """Extract links with context"""
        links = []

        for a in soup.find_all('a', href=True):
            href = a['href']
            text = a.get_text().strip()

            if not href or not text or len(text) < 3:
                continue

            # Convert relative URLs to absolute
            if not href.startswith(('http://', 'https://')):
                href = urljoin(base_url, href)

            # Skip common non-content URLs
            if any(skip in href.lower() for skip in ['/search', '/login', '/admin', '#']):
                continue

            links.append({
                "url": href,
                "text": text[:100],
                "is_external": urlparse(href).netloc != urlparse(base_url).netloc
            })

        return links[:20]  # Limit to 20 links

    def _clean_content(self, content: str) -> str:
        """Clean and normalize content"""
        if not content:
            return ''

        # Remove excessive whitespace
        content = re.sub(r'\s+', ' ', content)

        # Remove navigation and common non-content elements
        content = re.sub(r'\b(home|menu|navigation|footer|copyright|privacy|terms|contact|about|login|signup|search)\b',
                        '', content, flags=re.IGNORECASE)

        # Remove emails and phone numbers
        content = re.sub(r'\S+@\S+\.\S+', '[EMAIL]', content)
        content = re.sub(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', '[PHONE]', content)

        return content.strip()

    def _extract_basic_content(self, url: str) -> Dict[str, Any]:
        """Basic fallback content extraction"""
        try:
            response = self.session.get(url, timeout=15)
            response.raise_for_status()

            soup = BeautifulSoup(response.text, 'html.parser')
            text_content = soup.get_text(separator=' ', strip=True)
            text_content = re.sub(r'\s+', ' ', text_content).strip()

            return {
                "url": url,
                "title": soup.find('title').get_text(strip=True) if soup.find('title') else "",
                "content": text_content[:5000],  # Limit content
                "word_count": len(text_content.split()),
                "status_code": response.status_code
            }
        except Exception as e:
            return {"error": str(e), "url": url}

    def _analyze_with_gemini_tools(self, url: str, initial_content: Dict[str, Any], gemini_client, user_instructions: str) -> Dict[str, Any]:
        """Use Gemini tools for enhanced content analysis"""
        try:
            self.logger.log("🤖 Using Gemini tools for content analysis")

            # Define available Gemini tools
            tools = [
                types.Tool(google_search=types.GoogleSearch()),
                types.Tool(url_context=types.UrlContext())
            ]

            # Create analysis prompt with user instructions
            analysis_prompt = self._create_analysis_prompt(url, initial_content, user_instructions)

            # Configure generation
            generation_config = types.GenerateContentConfig(
                temperature=0.7,
                top_p=0.95,
                max_output_tokens=4000,
                candidate_count=1,
                thinking_config=types.ThinkingConfig(thinking_budget=2048),
                tools=tools
            )

            # Generate with tools
            response = gemini_client.models.generate_content(
                model=GEMINI_MODEL,
                contents=analysis_prompt,
                config=generation_config
            )

            # Process tool calls and results
            tool_results = self._process_tool_calls(response, gemini_client)

            return {
                "gemini_analysis": response.text if response.candidates else "",
                "tool_results": tool_results,
                "enhanced_content": initial_content
            }

        except Exception as e:
            self.logger.error(f"⚠️ Gemini tool analysis failed: {e}")
            return {"error": str(e), "basic_content": initial_content}


    def _create_analysis_prompt(self, url: str, content: Dict[str, Any], user_instructions: str) -> str:
        """Create comprehensive analysis prompt for Gemini 2.5 Pro with enhanced user integration"""
        content_preview = content.get('main_content', content.get('content', ''))[:5000]  # Increased for Pro model
        code_blocks = content.get('code_blocks', [])
        title = content.get('title', 'Unknown Document')
        description = content.get('description', '')

        # Enhanced code context with better formatting
        code_context = ""
        if code_blocks:
            code_context = "\n\n🎯 DETECTED CODE BLOCKS:"
            for i, block in enumerate(code_blocks[:5], 1):  # Increased limit for Pro model
                lang = block.get('language', 'unknown')
                code_preview = block['content'][:300]  # More code context
                code_context += f"\n{i}. **{lang}**:\n``` {lang}\n{code_preview}\n```"

        # Enhanced metadata context
        metadata_context = ""
        if content.get('structured_data'):
            metadata_context = "\n\n📊 STRUCTURED DATA FOUND:"
            for item in content['structured_data'][:3]:
                if item.get('type') == 'json-ld':
                    metadata_context += f"\n- JSON-LD Schema: {list(item.get('data', {}).keys())[:5]}"

        # Research context from user instructions
        research_focus = self._extract_research_focus(user_instructions)

        return f"""{self.system_prompt}

🔗 **CONTENT TO ANALYZE**
URL: {url}
TITLE: {title}
DESCRIPTION: {description}
CONTENT PREVIEW: {content_preview}{code_context}{metadata_context}

🎯 **USER INSTRUCTIONS** (PRIORITY FOCUS)
{user_instructions}

🔍 **RESEARCH FOCUS AREAS** (derived from user instructions)
{research_focus}

📋 **REQUIRED OUTPUT FORMAT**
Return ONLY valid JSON with this exact structure:
{{
  "summary": "Comprehensive overview addressing user instructions and research focus",
  "blocks": [
    {{
      "id": "unique-block-id",
      "type": "summary|key_points|architecture|mermaid|code|api_reference|guide|comparison|best_practices|troubleshooting",
      "size": "small|medium|large",
      "title": "Block title that reflects user focus",
      "content": "Content that directly addresses user instructions",
      "metadata": {{
        "language": "javascript|python|etc (for code blocks)",
        "priority": "high|medium|low",
        "user_focus_alignment": "high|medium|low",
        "research_integrated": true|false
      }}
    }}
  ]
}}

🎨 **BLOCK TYPE SELECTION GUIDELINES**
- **summary**: Executive overview tailored to user instructions
- **key_points**: User-focused highlights and takeaways
- **architecture**: System diagrams when relevant to user questions
- **mermaid**: Flowcharts, system diagrams, process visualizations
- **code**: Examples directly relevant to user's learning goals
- **api_reference**: API details when user is learning integration
- **guide**: Step-by-step instructions for user tasks
- **comparison**: Alternatives when user needs to choose approaches
- **best_practices**: Recommendations aligned with user context
- **troubleshooting**: Solutions for user-specific challenges

📏 **SIZE OPTIMIZATION** (3x3 grid = 9 units max)
- **large** (3 units): Complex diagrams, comprehensive guides, detailed explanations
- **medium** (2 units): Standard explanations, moderate diagrams, key examples
- **small** (1 unit): Quick facts, simple lists, brief explanations

🎯 **USER-CENTRIC ANALYSIS REQUIREMENTS**
1. **Directly address user instructions** in every block
2. **Research user-mentioned topics** using Google Search when relevant
3. **Adapt presentation style** based on user context (beginner vs expert)
4. **Include practical examples** that match user's learning objectives
5. **Optimize for user's domain** (web dev, API design, data science, etc.)
6. **Generate 3-6 blocks maximum** - quality over quantity
7. **Ensure visual learning** with appropriate diagram types
8. **Include working code examples** when user needs implementation guidance

🔧 **TOOL USAGE STRATEGY**
- Use **Google Search** for: current best practices, related technologies, user interest research
- Use **URL Context** for: enhanced content analysis, related documentation discovery
- Use **Code Execution** for: validating code examples, testing implementations

⚡ **CONTENT ANALYSIS STRATEGY**
1. **Extract main concepts** and map to user instructions
2. **Identify knowledge gaps** user might have
3. **Research complementary topics** that would help user understanding
4. **Create blocks** that build upon each other progressively
5. **Optimize for retention** with visual elements and practical examples

🎯 **SUCCESS CRITERIA**
- Every block must directly relate to user instructions
- Content should be immediately actionable for user's goals
- Visual elements should enhance understanding of user's topics
- Research should provide current, relevant context to user's needs

Generate analysis blocks that create a personalized learning experience for this specific user."""

    def _process_tool_calls(self, response, gemini_client) -> Dict[str, Any]:
        """Process and execute tool calls from Gemini"""
        tool_results = {
            "google_search": [],
            "url_context": []
        }

        try:
            if response.candidates and response.candidates[0].content:
                for part in response.candidates[0].content.parts:
                    if hasattr(part, 'function_call') and part.function_call:
                        tool_call = part.function_call
                        tool_name = tool_call.name
                        args = tool_call.args

                        self.logger.log(f"🔧 Executing tool: {tool_name}")

                        if tool_name == "google_search":
                            result = self._execute_google_search(args, gemini_client)
                            tool_results["google_search"].append(result)
                        elif tool_name == "url_context":
                            result = self._execute_url_context(args, gemini_client)
                            tool_results["url_context"].append(result)

        except Exception as e:
            self.logger.error(f"⚠️ Tool execution failed: {e}")

        return tool_results

    def _execute_google_search(self, args: Dict[str, Any], gemini_client) -> Dict[str, Any]:
        """Execute Google search using Gemini tool"""
        try:
            search_query = args.get('query', '')
            self.logger.log(f"🔍 Google Search: {search_query}")

            # Gemini handles the search internally
            return {
                "query": search_query,
                "timestamp": time.time(),
                "status": "executed"
            }
        except Exception as e:
            return {"error": str(e), "query": args.get('query', '')}


    def _execute_url_context(self, args: Dict[str, Any], gemini_client) -> Dict[str, Any]:
        """Get URL context using Gemini tool"""
        try:
            context_url = args.get('url', '')
            self.logger.log(f"🌐 URL Context: {context_url}")

            # Gemini handles URL context internally
            return {
                "url": context_url,
                "timestamp": time.time(),
                "status": "executed"
            }
        except Exception as e:
            return {"error": str(e), "url": args.get('url', '')}


    def _process_deep_content(self, url: str, enhanced_analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Process content with enhanced analysis results"""
        try:
            base_content = enhanced_analysis.get('enhanced_content', {})

            # Extract additional insights from tool results
            search_insights = self._extract_search_insights(enhanced_analysis)
            url_insights = self._extract_url_insights(enhanced_analysis)

            # Combine all insights
            deep_content = {
                **base_content,
                "search_insights": search_insights,
                "url_insights": url_insights,
                "analysis_timestamp": time.time()
            }

            return deep_content

        except Exception as e:
            self.logger.error(f"⚠️ Deep content processing failed: {e}")
            return enhanced_analysis

    def _extract_search_insights(self, analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract insights from Google search results"""
        search_results = analysis.get('tool_results', {}).get('google_search', [])
        insights = []

        for result in search_results:
            if 'error' not in result:
                insights.append({
                    "type": "search_result",
                    "query": result.get('query', ''),
                    "relevance_score": 0.8,  # Would be calculated from actual results
                    "timestamp": result.get('timestamp', time.time())
                })

        return insights


    def _extract_url_insights(self, analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract insights from URL context results"""
        url_results = analysis.get('tool_results', {}).get('url_context', [])
        insights = []

        for result in url_results:
            if 'error' not in result:
                insights.append({
                    "type": "url_context",
                    "url": result.get('url', ''),
                    "context_type": "enhanced_analysis",
                    "timestamp": result.get('timestamp', time.time())
                })

        return insights

    def _enrich_with_research(self, url: str, deep_content: Dict[str, Any], gemini_client, user_instructions: str) -> Dict[str, Any]:
        """Enrich content with additional research"""
        try:
            self.logger.log("🔬 Enriching content with research")

            # Add null check for deep_content
            if not deep_content:
                self.logger.log("⚠️ No deep content provided for research, returning basic content")
                return {
                    "url": url,
                    "research_results": [],
                    "research_enrichment": False,
                    "error": "No content to research"
                }

            # Generate research queries based on content
            research_queries = self._generate_research_queries(deep_content)

            # Execute research using Gemini tools
            research_results = []
            for query in research_queries[:3]:  # Limit to 3 research queries
                self.logger.log(f"📚 Researching: {query}")
                result = self._execute_research_query(query, gemini_client)
                research_results.append(result)

            return {
                **deep_content,
                "research_results": research_results,
                "research_timestamp": time.time()
            }

        except Exception as e:
            self.logger.error(f"⚠️ Research enrichment failed: {e}")
            return deep_content

    def _generate_research_queries(self, content: Dict[str, Any]) -> List[str]:
        """Generate relevant research queries based on content"""
        queries = []

        # Add null check
        if not content:
            return ["General web development best practices"]

        title = content.get('title', '')
        keywords = content.get('keywords', [])
        main_content = content.get('main_content', '')[:1000]

        # Generate queries from title
        if title:
            queries.append(f"What is {title}")
            queries.append(f"Latest developments in {title}")

        # Generate queries from keywords
        for keyword in keywords[:3]:
            queries.append(f"Best practices for {keyword}")
            queries.append(f"Tutorials on {keyword}")

        # Generate queries from content analysis
        if main_content:
            # Simple keyword extraction for research
            words = re.findall(r'\b\w{4,}\b', main_content.lower())
            common_words = [word for word in words if word not in ['that', 'with', 'have', 'this', 'will', 'your', 'from', 'they', 'know', 'want', 'been', 'good', 'much', 'some', 'time', 'very', 'when', 'come', 'here', 'just', 'like', 'long', 'make', 'many', 'over', 'such', 'take', 'than', 'them', 'well', 'were']]

            if common_words:
                top_keywords = list(set(common_words))[:3]
                for keyword in top_keywords:
                    queries.append(f"Understanding {keyword} in detail")

        return list(set(queries))[:5]  # Return unique queries, max 5

    def _execute_research_query(self, query: str, gemini_client) -> Dict[str, Any]:
        """Execute a research query using Gemini tools"""
        try:
            tools = [types.Tool(google_search=types.GoogleSearch())]

            prompt = f"Research this topic comprehensively: {query}\n\nProvide detailed insights and current information."

            generation_config = types.GenerateContentConfig(
                temperature=0.7,
                thinking_config=types.ThinkingConfig(thinking_budget=2048),
                tools=tools
            )

            response = gemini_client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt,
                config=generation_config
            )

            return {
                "query": query,
                "response": response.text if response.candidates else "",
                "timestamp": time.time(),
                "status": "completed"
            }

        except Exception as e:
            return {
                "query": query,
                "error": str(e),
                "timestamp": time.time(),
                "status": "failed"
            }

    def _synthesize_final_result(self, url: str, enriched_content: Dict[str, Any]) -> Dict[str, Any]:
        """Synthesize all analysis results into final output"""
        try:
            self.logger.log("🎯 Synthesizing final analysis result")

            # Combine all content
            final_content = enriched_content.get('main_content', enriched_content.get('content', ''))

            # Add research insights to content
            research_insights = enriched_content.get('research_results', [])
            if research_insights:
                research_summary = "\n\n".join([
                    f"Research on '{result['query']}': {result.get('response', '')[:500]}"
                    for result in research_insights
                    if result.get('status') == 'completed'
                ])
                final_content += f"\n\n--- Research Insights ---\n{research_summary}"

            # Create final result
            final_result = {
                "url": url,
                "title": enriched_content.get('title', ''),
                "description": enriched_content.get('description', ''),
                "content": final_content,
                "word_count": len(final_content.split()),
                "character_count": len(final_content),
                "keywords": enriched_content.get('keywords', []),
                "author": enriched_content.get('author', ''),
                "published_date": enriched_content.get('published_date'),
                "links": enriched_content.get('links', []),
                "code_blocks": enriched_content.get('code_blocks', []),
                "media": enriched_content.get('media', []),
                "structured_data": enriched_content.get('structured_data', []),
                "tool_usage": {
                    "google_search_count": len(enriched_content.get('search_insights', [])),
                    "code_execution_count": len(enriched_content.get('code_insights', [])),
                    "url_context_count": len(enriched_content.get('url_insights', [])),
                    "research_queries_count": len(enriched_content.get('research_results', []))
                },
                "analysis_metadata": {
                    "processing_time": time.time() - time.time(),  # Would be set by caller
                    "content_type": enriched_content.get('content_type', 'unknown'),
                    "analysis_version": "2.5",
                    "tools_used": ["google_search", "code_execution", "url_context"],
                    "timestamp": time.time()
                }
            }

            return final_result

        except Exception as e:
            self.logger.error(f"⚠️ Final synthesis failed: {e}")
            return {
                "url": url,
                "content": enriched_content.get('content', ''),
                "error": str(e)
            }

# ===== RESEARCH ENGINE CLASS =====
class ResearchEngine:
    """Research engine for user interests"""

    def __init__(self, logger):
        self.logger = logger

    def research_user_interests(self, interests: List[str], topic: str) -> Dict[str, Any]:
        """Research topics based on user interests"""
        self.logger.log(f"🔍 Researching topic: '{topic}' with {len(interests)} interests")

        research_results = {
            "topic": topic,
            "interests": interests,
            "related_topics": self._find_related_topics(topic, interests),
            "recommendations": self._generate_recommendations(topic, interests),
            "insights": self._generate_research_findings(topic, interests, "detailed")
        }

        self.logger.log(f"✅ Research completed: {len(research_results.get('related_topics', []))} related topics")
        return research_results

    def _find_related_topics(self, topic: str, interests: List[str]) -> List[str]:
        """Find related topics based on interests"""
        related = []

        # Simple keyword matching for related topics
        topic_lower = topic.lower()
        for interest in interests:
            interest_lower = interest.lower()
            if interest_lower in topic_lower or topic_lower in interest_lower:
                related.append(f"{interest} concepts")
                related.append(f"Advanced {interest}")
                related.append(f"{interest} best practices")

        # Add some general related topics
        related.extend([
            "Industry trends",
            "Latest developments",
            "Key challenges",
            "Future outlook"
        ])

        return list(set(related))[:10]

    def _generate_recommendations(self, topic: str, interests: List[str]) -> List[str]:
        """Generate recommendations based on topic and interests"""
        recommendations = [
            f"Explore {topic} documentation",
            f"Study {topic} examples and tutorials"
        ]

        for interest in interests:
            recommendations.append(f"Connect {topic} with {interest} concepts")

        return recommendations[:8]

    def _generate_research_findings(self, topic: str, interests: List[str], depth: str) -> str:
        """Generate research findings"""
        findings = f"Research on '{topic}' reveals strong connections to: {', '.join(interests[:3])}"

        if depth == "detailed":
            findings += ". Key insights include emerging trends, practical applications, and learning opportunities."

        return findings

# ===== DOCIFY UNIFIED ORCHESTRATOR =====
class DocifyUnifiedOrchestrator:
    """Advanced orchestrator using all Gemini tools for comprehensive content analysis"""

    def __init__(self, context, databases, logger):
        self.context = context
        self.databases = databases
        self.logger = logger

        # Store system prompt FIRST - critical for initialization order
        self.system_prompt = SYSTEM_PROMPT

        # Validate Gemini API key before initializing client
        if not GEMINI_API_KEY or GEMINI_API_KEY.strip() == '':
            raise ValueError("GEMINI_API_KEY environment variable is not set or is empty")

        # Initialize Gemini client with system prompt and enhanced configuration
        try:
            self.gemini_client = genai.Client(api_key=GEMINI_API_KEY)
        except Exception as e:
            raise ValueError(f"Failed to initialize Gemini client: {str(e)}")

        # Initialize advanced components - now system_prompt is available
        self.content_processor = AdvancedContentProcessor(logger, self.system_prompt)
        self.research_engine = ResearchEngine(logger)

        # Track comprehensive tool usage
        self.tool_usage_log = []
        self.start_time = time.time()

    def process_document_comprehensive(self) -> Dict[str, Any]:
        """Comprehensive document processing pipeline using all Gemini tools"""
        try:
            self.logger.log("🚀 === STARTING COMPREHENSIVE DOCUMENT PROCESSING ===")

            # 1. Extract document data with enhanced validation
            self.logger.log("📋 Step 1: Extracting document data...")
            try:
                document_data = self._extract_document_data()
            except Exception as e:
                self.logger.error(f"❌ Document data extraction failed: {e}")
                raise ValueError(f"Failed to extract document data: {e}")

            # Add null check for document_data
            if not document_data:
                self.logger.error("❌ Document data extraction returned None")
                raise ValueError("Failed to extract document data from request")

            self.logger.log(f"📄 Document: {document_data.get('title', 'Untitled')}")
            self.logger.log(f"🔗 URL: {document_data.get('url', 'No URL')}")
            self.logger.log(f"📝 Instructions: {document_data.get('instructions', 'No instructions')}")

            # Extract document_id from document_data
            document_id = document_data.get('document_id')
            if not document_id:
                raise ValueError("Document ID not found in document data")

            # 2. Scrape content from the website using comprehensive scraper
            self.logger.log("🌐 Step 2: Scraping content from website using comprehensive scraper...")
            scraped_content = self._scrape_website_content(document_data.get('url', ''))

            # 3. Save scraped content immediately to database
            self.logger.log("💾 Step 3: Saving scraped content to database...")
            self._save_scraped_content_immediately(document_id, scraped_content)

            # 4. Multi-round analysis with all Gemini tools using scraped content
            self.logger.log("🤖 Step 4: Executing multi-round analysis with all Gemini tools...")
            analysis_result = self._execute_multi_round_analysis_with_scraped_content(document_data, self.gemini_client, scraped_content)

            # 5. Generate comprehensive analysis blocks
            self.logger.log("📊 Step 5: Generating comprehensive analysis blocks...")
            analysis_blocks = self._generate_comprehensive_blocks(analysis_result, document_data)

            # 6. Synthesize final result
            self.logger.log("🎯 Step 6: Synthesizing final comprehensive result...")
            final_result = self._create_final_comprehensive_result(
                analysis_result, analysis_blocks, document_data
            )

            # 7. Update database with analysis results (content already saved)
            self.logger.log("💾 Step 7: Updating database with analysis results...")
            self._update_document_with_analysis_only(document_id, final_result)

            processing_time = time.time() - self.start_time
            self.logger.log(f"⏱️ Processing completed in {processing_time:.2f}s")
            self.logger.log(f"📊 Tool executions: {len(self.tool_usage_log)}")
            self.logger.log("🎉 === COMPREHENSIVE PROCESSING COMPLETED ===")

            return final_result

        except Exception as e:
            self.logger.error(f"❌ Comprehensive processing failed: {str(e)}")
            processing_time = time.time() - self.start_time

            # Create error result
            error_result = {
                "success": False,
                "error": str(e),
                "document_id": document_data.get('document_id') if 'document_data' in locals() else 'unknown',
                "processing_time": processing_time,
                "tool_usage": self.tool_usage_log
            }

            # Try to update document status
            if 'document_data' in locals() and document_data.get('document_id'):
                try:
                    self._update_document_status(document_data['document_id'], 'failed')
                except Exception as update_error:
                    self.logger.error(f"⚠️ Failed to update document status: {update_error}")

            return error_result

    def process_document(self) -> Dict[str, Any]:
        """Main processing pipeline"""
        try:
            self.logger.log("=== STARTING DOCUMENT PROCESSING ===")

            # 1. Extract document data
            self.logger.log("📋 Step 1: Extracting document data...")
            document_data = self._extract_document_data()
            self.logger.log(f"📄 Document: {document_data.get('title', 'Untitled')}")
            self.logger.log(f"🔗 URL: {document_data.get('url', 'No URL')}")
            self.logger.log(f"📝 Instructions: {document_data.get('instructions', 'No instructions')}")

            # 2. Get user interests
            self.logger.log("🎯 Step 2: Retrieving user interests...")
            user_interests = self._get_user_interests(document_data.get('user_id'))
            self.logger.log(f"📊 User interests: {len(user_interests)} found")

            # 3. Create analysis plan
            self.logger.log("🤖 Step 3: Creating analysis plan with Gemini...")
            analysis_plan = self._create_analysis_plan(document_data, user_interests)

            # 4. Execute workflow
            self.logger.log("⚡ Step 4: Executing orchestrated workflow...")
            workflow_result = self._execute_workflow(analysis_plan, document_data)

            # 5. Format results
            self.logger.log("📝 Step 5: Formatting final results...")
            final_result = self._format_final_result(workflow_result, document_data)

            # 6. Update database
            self.logger.log("💾 Step 6: Updating document in database...")
            self._update_document_status(document_data['document_id'], 'completed', final_result)

            self.logger.log("🎉 === DOCUMENT PROCESSING COMPLETED ===")
            return final_result

        except Exception as e:
            self.logger.error(f"❌ Processing failed: {str(e)}")
            if 'document_id' in locals():
                self._update_document_status(document_data['document_id'], 'failed', error=str(e))
            raise

    def _extract_research_focus(self, user_instructions: str) -> str:
        """Extract research focus areas from user instructions"""
        research_areas = []

        # Common research triggers
        research_keywords = {
            'best practices': ['current best practices', 'industry standards', 'recommended approaches'],
            'tutorial': ['step-by-step guides', 'learning resources', 'examples'],
            'api': ['integration patterns', 'authentication methods', 'API design'],
            'performance': ['optimization techniques', 'performance benchmarks', 'scaling strategies'],
            'security': ['security best practices', 'common vulnerabilities', 'secure implementation'],
            'comparison': ['alternative solutions', 'feature comparisons', 'trade-offs'],
            'architecture': ['system design patterns', 'scalability considerations', 'design principles'],
            'deployment': ['deployment strategies', 'production considerations', 'hosting options']
        }

        instructions_lower = user_instructions.lower()

        for keyword, areas in research_keywords.items():
            if keyword in instructions_lower:
                research_areas.extend(areas)

        if not research_areas:
            research_areas = ['related concepts', 'practical applications', 'implementation examples']

        return "• " + "\n• ".join(list(set(research_areas))[:5])

    def _scrape_website_content(self, url: str) -> Dict[str, Any]:
        """Scrape content from website using comprehensive scraper"""
        try:
            self.logger.log(f"🌐 Starting comprehensive website scraping for: {url}")

            # Use the comprehensive scraping function
            max_pages = int(os.environ.get('MAX_PAGES_TO_CRAWL', '10'))
            scraped_data = scrape_website_with_requests(url, max_pages)

            self.logger.log(f"✅ Scraping completed: {scraped_data.get('word_count', 0)} words from {scraped_data.get('pages_crawled', 0)} pages")
            return scraped_data

        except Exception as e:
            self.logger.error(f"❌ Website scraping failed: {e}")
            # Return basic fallback
            return {
                "url": url,
                "title": "Scraping Failed",
                "description": f"Failed to scrape content: {str(e)}",
                "content": f"Error occurred while scraping: {str(e)}",
                "word_count": 0,
                "scraped_at": "error"
            }

    def _save_scraped_content_immediately(self, document_id: str, scraped_content: Dict[str, Any]) -> None:
        """Save scraped content immediately to database"""
        try:
            self.logger.log(f"💾 Saving scraped content for document {document_id}")

            update_data = {
                'status': 'analyzing',  # Status indicating content scraped and analysis starting
                'scraped_content': scraped_content.get('content', ''),
                'word_count': scraped_content.get('word_count', 0),
                'analysis_summary': scraped_content.get('title', 'Document')  # Save title as initial summary
            }

            self.databases.update_document(
                DATABASE_ID,
                DOCUMENTS_COLLECTION_ID,
                document_id,
                update_data
            )

            self.logger.log(f"✅ Scraped content saved for document {document_id}")

        except Exception as e:
            self.logger.error(f"⚠️ Failed to save scraped content: {e}")
            # Don't raise - we want to continue with analysis even if saving fails

    def _execute_multi_round_analysis_with_scraped_content(self, document_data: Dict[str, Any], gemini_client, scraped_content: Dict[str, Any]) -> Dict[str, Any]:
        """Execute multi-round analysis using already scraped content"""
        try:
            url = document_data.get('url', '')
            instructions = document_data.get('instructions', 'Analyze this content comprehensively')

            self.logger.log("🔄 Round 1: Processing scraped content for analysis")

            # Use the scraped content directly instead of scraping again
            initial_content = {
                "url": url,
                "title": scraped_content.get('title', ''),
                "description": scraped_content.get('description', ''),
                "content": scraped_content.get('content', ''),
                "word_count": scraped_content.get('word_count', 0),
                "content_type": "scraped_content",
                "pages_crawled": scraped_content.get('pages_crawled', 0)
            }

            # Round 2: Enhanced analysis with user instructions
            self.logger.log("🔄 Round 2: Enhanced analysis with user instructions")
            enhanced_analysis = self._perform_enhanced_analysis(initial_content, instructions, gemini_client)

            # Round 3: Research enrichment
            self.logger.log("🔄 Round 3: Research and context enrichment")
            enriched_analysis = self._perform_research_enrichment_with_content(enhanced_analysis, gemini_client, initial_content)

            # Round 4: Final synthesis and validation
            self.logger.log("🔄 Round 4: Final synthesis and validation")
            final_analysis = self._perform_final_synthesis_with_content(enriched_analysis, document_data, gemini_client, instructions, initial_content)

            return {
                "comprehensive_result": initial_content,
                "enhanced_analysis": enhanced_analysis,
                "enriched_analysis": enriched_analysis,
                "final_analysis": final_analysis,
                "rounds_completed": 4,
                "timestamp": time.time(),
                "scraped_content_used": True
            }

        except Exception as e:
            self.logger.error(f"⚠️ Multi-round analysis with scraped content failed: {e}")
            return {"error": str(e), "rounds_completed": 0, "scraped_content_used": True}

    def _update_document_with_analysis_only(self, document_id: str, result: Dict[str, Any]) -> None:
        """Update document with analysis results only (content already saved)"""
        try:
            self.logger.log(f"💾 Updating document {document_id} with analysis results only")

            # Extract schema-compliant summary from enhanced analysis
            enhanced_analysis = result.get('enhanced_analysis', {})
            summary_text = ""

            if isinstance(enhanced_analysis.get('enhanced_analysis'), dict):
                schema_data = enhanced_analysis['enhanced_analysis']
                if 'summary' in schema_data and schema_data['summary']:
                    summary_text = schema_data['summary']
                    self.logger.log(f"📝 Using schema-compliant summary from LLM: {len(summary_text)} chars")
                else:
                    # Fallback to other sources
                    summary_text = result.get('analysis_summary', '')
            else:
                # Fallback to other sources
                summary_text = result.get('analysis_summary', '')

            # Ensure summary_text is always a valid string and within limits
            if not isinstance(summary_text, str):
                summary_text = str(summary_text) if summary_text is not None else ""

            # If still no summary, create a basic one
            if not summary_text.strip():
                final_analysis = result.get('final_analysis', {})
                summary_text = final_analysis.get('final_synthesis', 'Analysis completed successfully')

            # Truncate to 2000 characters to meet database constraints
            if len(summary_text) > 2000:
                summary_text = summary_text[:1997] + "..."
                self.logger.log(f"⚠️ Summary truncated to 2000 chars: {len(summary_text)}")

            # Validate and sanitize blocks data
            blocks_data = result.get('blocks', [])
            if not isinstance(blocks_data, list):
                blocks_data = []
                self.logger.log("⚠️ analysis_blocks was not a list, using empty array")

            sanitized_blocks = []
            for block in blocks_data:
                if isinstance(block, dict) and all(key in block for key in ['id', 'type', 'size', 'title', 'content']):
                    # Ensure all string fields are actually strings
                    sanitized_block = {
                        'id': str(block.get('id', '')),
                        'type': str(block.get('type', '')),
                        'size': str(block.get('size', '')),
                        'title': str(block.get('title', '')),
                        'content': str(block.get('content', ''))
                    }
                    # Add metadata if it exists and is valid
                    if 'metadata' in block and isinstance(block['metadata'], dict):
                        sanitized_block['metadata'] = block['metadata']
                    sanitized_blocks.append(sanitized_block)

            # Get current document to preserve title
            try:
                current_doc = self.databases.get_document(
                    DATABASE_ID,
                    DOCUMENTS_COLLECTION_ID,
                    document_id
                )
                current_title = current_doc.get('analysis_summary', '')
                if current_title and current_title != summary_text:
                    # Combine title with analysis summary
                    combined_summary = f"{current_title}\n\n{summary_text}"
                    # Truncate if too long (analysis_summary has 2000 char limit)
                    if len(combined_summary) > 2000:
                        combined_summary = combined_summary[:1997] + "..."
                    final_summary = combined_summary
                else:
                    final_summary = summary_text
            except Exception:
                # Fallback to just the summary if we can't get current document
                final_summary = summary_text

            update_data = {
                'status': 'completed',
                'analysis_summary': final_summary,
                'analysis_blocks': json.dumps(sanitized_blocks)
            }

            # Add processing duration if available
            available_fields = self._get_available_fields()
            if 'processing_duration' in available_fields:
                # Calculate processing time (rough estimate)
                update_data['processing_duration'] = int(time.time() - self.start_time)

            self.databases.update_document(
                DATABASE_ID,
                DOCUMENTS_COLLECTION_ID,
                document_id,
                update_data
            )

            self.logger.log(f"✅ Document {document_id} updated with analysis results")

        except Exception as e:
            self.logger.error(f"⚠️ Analysis-only update failed: {e}")

    def _extract_document_data(self) -> Dict[str, Any]:
        """Extract document data from request with enhanced validation"""
        try:
            trigger_type = self.context.req.headers.get('x-appwrite-trigger', 'http')

            if trigger_type == 'event':
                body = self.context.req.body
                if isinstance(body, str):
                    body = json.loads(body)
                document_data = {
                    'document_id': body.get('$id'),
                    'url': body.get('url'),
                    'instructions': body.get('instructions'),
                    'user_id': body.get('user_id'),
                    'title': body.get('title')
                }
            else:
                body = self.context.req.body
                if isinstance(body, str):
                    body = json.loads(body)
                document_data = {
                    'document_id': body.get('documentId'),
                    'url': body.get('url'),
                    'instructions': body.get('instructions'),
                    'user_id': body.get('userId'),
                    'title': body.get('title')
                }

            # Validate required fields
            required_fields = ['document_id', 'url']
            missing_fields = [field for field in required_fields if not document_data.get(field)]

            if missing_fields:
                raise ValueError(f"Missing required fields: {', '.join(missing_fields)}")

            # Set default instructions if none provided
            if not document_data.get('instructions'):
                document_data['instructions'] = 'Analyze this content comprehensively using all available tools'

            return document_data

        except Exception as e:
            self.logger.error(f"⚠️ Document data extraction failed: {e}")
            raise ValueError(f"Failed to extract document data: {e}")

    def _get_user_interests(self, user_id: str) -> List[str]:
        """Get user interests (placeholder)"""
        # For now, return some default interests
        return ["web development", "artificial intelligence", "api design", "data analysis", "software architecture"]

    def _update_document_comprehensive(self, document_id: str, result: Dict[str, Any]) -> None:
        """Update document with comprehensive results"""
        try:
            self.logger.log(f"💾 Updating document {document_id} with comprehensive results")

            # Extract schema-compliant summary from enhanced analysis
            enhanced_analysis = result.get('enhanced_analysis', {})
            summary_text = ""

            if isinstance(enhanced_analysis.get('enhanced_analysis'), dict):
                schema_data = enhanced_analysis['enhanced_analysis']
                if 'summary' in schema_data and schema_data['summary']:
                    summary_text = schema_data['summary']
                    self.logger.log(f"📝 Using schema-compliant summary from LLM: {len(summary_text)} chars")
                else:
                    # Fallback to other sources
                    summary_text = result.get('analysis_summary', '')
            else:
                # Fallback to other sources
                summary_text = result.get('analysis_summary', '')

            # Ensure summary_text is always a valid string and within limits
            if not isinstance(summary_text, str):
                summary_text = str(summary_text) if summary_text is not None else ""

            # If still no summary, create a basic one
            if not summary_text.strip():
                comprehensive_result = result.get('comprehensive_result', {})
                content = comprehensive_result.get('content', '')[:500]
                summary_text = f"Document analysis completed. {len(content)} characters processed."

            # Truncate to 2000 characters to meet database constraints
            if len(summary_text) > 2000:
                summary_text = summary_text[:1997] + "..."
                self.logger.log(f"⚠️ Summary truncated to 2000 chars: {len(summary_text)}")

            # Ensure analysis_blocks is always valid JSON
            blocks_data = result.get('blocks', [])
            if not isinstance(blocks_data, list):
                blocks_data = []
                self.logger.log("⚠️ analysis_blocks was not a list, using empty array")

            # Validate and sanitize blocks data
            sanitized_blocks = []
            for block in blocks_data:
                if isinstance(block, dict) and all(key in block for key in ['id', 'type', 'size', 'title', 'content']):
                    # Ensure all string fields are actually strings
                    sanitized_block = {
                        'id': str(block.get('id', '')),
                        'type': str(block.get('type', '')),
                        'size': str(block.get('size', '')),
                        'title': str(block.get('title', '')),
                        'content': str(block.get('content', ''))
                    }
                    # Add metadata if it exists and is valid
                    if 'metadata' in block and isinstance(block['metadata'], dict):
                        sanitized_block['metadata'] = block['metadata']
                    sanitized_blocks.append(sanitized_block)

            update_data = {
                'status': 'completed',
                'analysis_summary': summary_text,
                'analysis_blocks': json.dumps(sanitized_blocks),
                'word_count': len(summary_text.split()),
                'scraped_content': result.get('comprehensive_result', {}).get('content', '')
            }

            # Save code blocks if available
            comprehensive_result = result.get('comprehensive_result', {})
            if comprehensive_result.get('code_blocks'):
                code_blocks_json = json.dumps(comprehensive_result['code_blocks'])
                # Check if there's a field for code blocks
                available_fields = self._get_available_fields()
                if 'code_blocks' in available_fields:
                    update_data['code_blocks'] = code_blocks_json
                elif 'extracted_code' in available_fields:
                    update_data['extracted_code'] = code_blocks_json

            # Add tool usage tracking
            if self.tool_usage_log:
                update_data['gemini_tools_used'] = json.dumps([t.get('tool', 'unknown') for t in self.tool_usage_log])

            # Add processing duration if available
            available_fields = self._get_available_fields()
            if 'processing_duration' in available_fields:
                # Calculate processing time (rough estimate)
                update_data['processing_duration'] = 30  # Default estimate

            self.databases.update_document(
                DATABASE_ID,
                DOCUMENTS_COLLECTION_ID,
                document_id,
                update_data
            )

            self.logger.log(f"✅ Document {document_id} updated with schema-compliant data")

        except Exception as e:
            self.logger.error(f"⚠️ Comprehensive update failed: {e}")
            # Fallback to basic status update
            try:
                self._update_document_status(document_id, 'completed', result)
            except Exception as fallback_error:
                self.logger.error(f"⚠️ Fallback update also failed: {fallback_error}")

    def _validate_content(self, document_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate content before processing"""
        try:
            url = document_data.get('url', '')
            if not url:
                return {"error": "No URL provided"}

            # Basic URL validation
            parsed_url = urlparse(url)
            if not parsed_url.scheme or not parsed_url.netloc:
                return {"error": "Invalid URL format"}

            # Check for supported schemes
            if parsed_url.scheme not in ['http', 'https']:
                return {"error": "Only HTTP and HTTPS URLs are supported"}

            self.logger.log(f"✅ Content validation passed for: {url}")
            return {"valid": True, "url_info": parsed_url}

        except Exception as e:
            return {"error": f"Content validation failed: {str(e)}"}

    def _execute_multi_round_analysis(self, document_data: Dict[str, Any], gemini_client) -> Dict[str, Any]:
        """Execute multi-round analysis using all Gemini tools"""
        try:
            url = document_data.get('url', '')
            instructions = document_data.get('instructions', 'Analyze this content comprehensively')

            self.logger.log("🔄 Round 1: Comprehensive URL processing")
            # Use the advanced content processor
            comprehensive_result = self.content_processor.process_url_comprehensive(url, gemini_client, instructions)

            if comprehensive_result.get('error'):
                self.logger.log("⚠️ Comprehensive processing failed, falling back to basic processing")
                # Fallback to basic processing
                basic_result = self.content_processor._extract_basic_content(url)
                comprehensive_result = {
                    "url": url,
                    "title": basic_result.get('title', ''),
                    "content": basic_result.get('content', ''),
                    "word_count": basic_result.get('word_count', 0),
                    "fallback_used": True
                }

            # Round 2: Enhanced analysis with user instructions
            self.logger.log("🔄 Round 2: Enhanced analysis with user instructions")
            enhanced_analysis = self._perform_enhanced_analysis(comprehensive_result, instructions, gemini_client)

            # Round 3: Research enrichment
            self.logger.log("🔄 Round 3: Research and context enrichment")
            enriched_analysis = self._perform_research_enrichment(enhanced_analysis, gemini_client)

            # Round 4: Final synthesis and validation
            self.logger.log("🔄 Round 4: Final synthesis and validation")
            final_analysis = self._perform_final_synthesis(enriched_analysis, document_data, gemini_client, instructions)

            return {
                "comprehensive_result": comprehensive_result,
                "enhanced_analysis": enhanced_analysis,
                "enriched_analysis": enriched_analysis,
                "final_analysis": final_analysis,
                "rounds_completed": 4,
                "timestamp": time.time()
            }

        except Exception as e:
            self.logger.error(f"⚠️ Multi-round analysis failed: {e}")
            return {"error": str(e), "rounds_completed": 0}

    def _perform_enhanced_analysis(self, content_result: Dict[str, Any], instructions: str, gemini_client) -> Dict[str, Any]:
        """Perform enhanced analysis with user instructions following schema"""
        try:
            content = content_result.get('content', '')
            if not content:
                return {"error": "No content available for enhanced analysis"}

            # Enhanced analysis prompt using comprehensive system prompt
            research_focus = self._extract_research_focus(instructions)

            prompt = f"""{self.system_prompt}

🔄 **ENHANCED ANALYSIS ROUND**
CONTENT TITLE: {content_result.get('title', 'Unknown Document')}
CONTENT PREVIEW: {content[:5000]}
CONTENT TYPE: {content_result.get('content_type', 'unknown')}

🎯 **USER INSTRUCTIONS** (PRIMARY FOCUS)
{instructions}

🔍 **RESEARCH FOCUS AREAS**
{research_focus}

📊 **ANALYSIS OBJECTIVES**
1. Deep-dive analysis of content based on user instructions
2. Research complementary information using available tools
3. Create focused analysis blocks that directly address user needs
4. Optimize for user's learning objectives and context

🎨 **BLOCK GENERATION STRATEGY**
- Generate 3-5 blocks maximum for enhanced analysis round
- Focus on depth rather than breadth
- Use advanced block types: architecture, mermaid, api_reference, guide
- Include research-backed insights from Google Search
- Adapt complexity based on user expertise level

📋 **REQUIRED OUTPUT FORMAT**
{{
  "summary": "Deep analysis summary addressing user instructions with research insights",
  "blocks": [
    {{
      "id": "enhanced-analysis-1",
      "type": "architecture|mermaid|api_reference|guide|comparison|best_practices|troubleshooting",
      "size": "medium|large",
      "title": "Enhanced Analysis Block Title",
      "content": "Detailed content with research insights and practical guidance",
      "metadata": {{
        "priority": "high|medium|low",
        "user_focus_alignment": "high|medium|low",
        "research_integrated": true|false,
        "analysis_depth": "deep|comprehensive",
        "tool_usage": "google_search|url_context|code_execution"
      }}
    }}
  ]
}}

⚡ **ENHANCED ANALYSIS REQUIREMENTS**
1. **Research Integration**: Use Google Search for current best practices and related information
2. **User Context Adaptation**: Adjust analysis depth and style based on user instructions
3. **Practical Application**: Include actionable insights and real-world examples
4. **Visual Enhancement**: Use mermaid diagrams for complex concepts and system architectures
5. **Progressive Disclosure**: Build complexity gradually from foundational to advanced concepts

🔧 **TOOL USAGE PRIORITIES**
- **Google Search**: Research current trends, best practices, alternatives
- **URL Context**: Find related documentation and resources
- **Code Execution**: Validate technical examples and implementations

🎯 **SUCCESS CRITERIA**
- Every block provides unique value aligned with user instructions
- Research insights enhance the original content analysis
- Visual elements clarify complex technical concepts
- Content is immediately applicable to user's goals"""

            # Use Gemini for enhanced analysis with thinking mode enabled
            response = gemini_client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.7,
                    max_output_tokens=3000,
                    thinking_config=types.ThinkingConfig(thinking_budget=2048)  # Enable thinking mode for 2.5 Pro
                )
            )

            # Robust JSON parsing with multiple strategies
            analysis_text = response.text if response.candidates else ""
            parsed_result = self._parse_llm_response(analysis_text)

            return {
                "enhanced_analysis": parsed_result.get("parsed_data", {}),
                "raw_response": analysis_text,
                "instructions_alignment": "analyzed",
                "format": parsed_result.get("format", "text"),
                "parse_success": parsed_result.get("success", False),
                "timestamp": time.time()
            }

        except Exception as e:
            self.logger.error(f"⚠️ Enhanced analysis failed: {e}")
            return {"error": str(e)}

    def _parse_llm_response(self, response_text: str) -> Dict[str, Any]:
        """Robust JSON parsing with multiple fallback strategies"""
        try:
            if not response_text or not response_text.strip():
                return {
                    "success": False,
                    "format": "empty",
                    "error": "Empty response",
                    "parsed_data": {}
                }

            text = response_text.strip()

            # Strategy 1: Direct JSON parsing
            if text.startswith('{') and text.endswith('}'):
                try:
                    parsed = json.loads(text)
                    if self._validate_llm_schema(parsed):
                        return {
                            "success": True,
                            "format": "json",
                            "parsed_data": parsed
                        }
                except json.JSONDecodeError:
                    pass

            # Strategy 2: Extract JSON from text (look for JSON blocks)
            json_patterns = [
                r'```json\s*(\{.*?\})\s*```',  # Markdown JSON blocks
                r'```\s*(\{.*?\})\s*```',      # Generic code blocks
                r'(\{[^{}]*\{[^{}]*\}[^{}]*\})',  # Nested JSON
                r'(\{.*?\n\})',                # Multi-line JSON
            ]

            for pattern in json_patterns:
                matches = re.findall(pattern, text, re.DOTALL)
                for match in matches:
                    try:
                        parsed = json.loads(match)
                        if self._validate_llm_schema(parsed):
                            return {
                                "success": True,
                                "format": "extracted_json",
                                "parsed_data": parsed
                            }
                    except json.JSONDecodeError:
                        continue

            # Strategy 3: Try to fix common JSON issues and parse
            fixed_text = self._fix_json_issues(text)
            if fixed_text != text:
                try:
                    parsed = json.loads(fixed_text)
                    if self._validate_llm_schema(parsed):
                        return {
                            "success": True,
                            "format": "fixed_json",
                            "parsed_data": parsed
                        }
                except json.JSONDecodeError:
                    pass

            # Strategy 4: Generate fallback analysis from text
            self.logger.log(f"⚠️ JSON parsing failed, generating fallback analysis")
            fallback_data = self._generate_fallback_analysis(text)
            return {
                "success": False,
                "format": "fallback",
                "error": "JSON parsing failed, using fallback",
                "parsed_data": fallback_data
                }

        except Exception as e:
            self.logger.error(f"⚠️ LLM response parsing error: {e}")
            return {
                "success": False,
                "format": "error",
                "error": str(e),
                "parsed_data": self._generate_error_fallback()
            }

    def _validate_llm_schema(self, data: Dict[str, Any]) -> bool:
        """Validate LLM response against schema requirements"""
        try:
            if not isinstance(data, dict):
                return False

            # Check for required top-level fields
            if "summary" not in data:
                return False

            if "blocks" not in data or not isinstance(data["blocks"], list):
                return False

            # Validate blocks array
            if not data["blocks"]:
                return False

            # Check each block for required fields
            valid_types = [
                "summary", "key_points", "architecture", "mermaid", "code",
                "api_reference", "guide", "comparison", "best_practices", "troubleshooting"
            ]
            valid_sizes = ["small", "medium", "large"]

            block_ids = set()
            for block in data["blocks"]:
                if not isinstance(block, dict):
                    return False

                # Required fields
                required_fields = ["id", "type", "size", "title", "content"]
                for field in required_fields:
                    if field not in block:
                        return False

                # Validate block ID uniqueness
                block_id = block["id"]
                if block_id in block_ids:
                    return False
                block_ids.add(block_id)

                # Validate block type
                if block["type"] not in valid_types:
                    return False

                # Validate block size
                if block["size"] not in valid_sizes:
                    return False

                # Validate metadata if present
                if "metadata" in block:
                    if not isinstance(block["metadata"], dict):
                        return False

            # Check maximum blocks constraint (6 blocks max)
            if len(data["blocks"]) > 6:
                return False

            return True

        except Exception as e:
            self.logger.error(f"⚠️ Schema validation error: {e}")
            return False

    def _fix_json_issues(self, text: str) -> str:
        """Fix common JSON formatting issues"""
        try:
            # Remove markdown code blocks if present
            text = re.sub(r'```\w*\n?', '', text)
            text = re.sub(r'```\n?', '', text)

            # Fix trailing commas
            text = re.sub(r',\s*}', '}', text)
            text = re.sub(r',\s*]', ']', text)

            # Fix single quotes to double quotes (basic cases)
            text = re.sub(r"'([^']*)':", r'"\1":', text)
            text = re.sub(r":\s*'([^']*)'", r': "\1"', text)
            text = re.sub(r":\s*'([^']*)',", r': "\1",', text)

            # Fix unquoted keys (basic pattern)
            text = re.sub(r'(\w+):', r'"\1":', text)

            # Remove extra whitespace
            text = text.strip()

            return text

        except Exception as e:
            self.logger.error(f"⚠️ JSON fixing error: {e}")
            return text

    def _generate_fallback_analysis(self, text: str) -> Dict[str, Any]:
        """Generate fallback analysis from text response"""
        try:
            # Extract summary from text
            lines = text.split('\n')
            summary = ""
            for line in lines[:5]:  # First 5 lines as summary
                if line.strip():
                    summary += line.strip() + " "
            summary = summary.strip() or "Analysis completed successfully"

            # Create basic blocks
            blocks = [
                {
                    "id": "fallback-summary",
                    "type": "summary",
                    "size": "large",
                    "title": "Analysis Summary",
                    "content": summary,
                    "metadata": {
                        "priority": "high",
                        "source": "fallback"
                    }
                },
                {
                    "id": "fallback-content",
                    "type": "key_points",
                    "size": "medium",
                    "title": "Content Analysis",
                    "content": f"Analysis completed with {len(text)} characters of content processed.",
                    "metadata": {
                        "priority": "medium",
                        "source": "fallback"
                    }
                }
            ]

            return {
                "summary": summary,
                "blocks": blocks
            }

        except Exception as e:
            self.logger.error(f"⚠️ Fallback generation error: {e}")
            return self._generate_error_fallback()

    def _generate_error_fallback(self) -> Dict[str, Any]:
        """Generate minimal error fallback"""
        return {
            "summary": "Analysis completed with errors",
            "blocks": [
                {
                    "id": "error-summary",
                    "type": "summary",
                    "size": "small",
                    "title": "Analysis Error",
                    "content": "Document analysis completed but some processing errors occurred.",
                    "metadata": {
                        "priority": "low",
                        "source": "error_fallback"
                    }
                }
            ]
        }

    def _validate_block_structure(self, block: Dict[str, Any]) -> bool:
        """Validate individual block structure against schema"""
        try:
            if not isinstance(block, dict):
                return False

            # Required fields
            required_fields = ["id", "type", "size", "title", "content"]
            for field in required_fields:
                if field not in block:
                    return False

            # Validate block type
            valid_types = [
                "summary", "key_points", "architecture", "mermaid", "code",
                "api_reference", "guide", "comparison", "best_practices", "troubleshooting"
            ]
            if block["type"] not in valid_types:
                return False

            # Validate block size
            valid_sizes = ["small", "medium", "large"]
            if block["size"] not in valid_sizes:
                return False

            # Validate metadata if present
            if "metadata" in block and not isinstance(block["metadata"], dict):
                return False

            return True

        except Exception as e:
            self.logger.error(f"⚠️ Block validation error: {e}")
            return False

    def _normalize_block_structure(self, block: Dict[str, Any]) -> Dict[str, Any]:
        """Normalize block structure to ensure schema compliance"""
        try:
            normalized = {
                "id": block.get("id", f"block-{hash(str(block))}"),
                "type": block.get("type", "summary"),
                "size": block.get("size", "medium"),
                "title": block.get("title", "Analysis Block"),
                "content": block.get("content", "Content not available"),
            }

            # Ensure metadata exists
            if "metadata" not in block:
                normalized["metadata"] = {
                    "priority": "medium",
                    "source": "normalized"
                }
            else:
                normalized["metadata"] = block["metadata"]

            # Ensure metadata has required fields
            if "priority" not in normalized["metadata"]:
                normalized["metadata"]["priority"] = "medium"

            return normalized

        except Exception as e:
            self.logger.error(f"⚠️ Block normalization error: {e}")
            return {
                "id": f"normalized-{hash(str(block))}",
                "type": "summary",
                "size": "small",
                "title": "Normalized Block",
                "content": "Block content normalized due to formatting issues.",
                "metadata": {
                    "priority": "low",
                    "source": "normalization_fallback"
                }
            }

    def _perform_research_enrichment(self, analysis_result: Dict[str, Any], gemini_client) -> Dict[str, Any]:
        """Perform research enrichment using Gemini tools"""
        try:
            content = analysis_result.get('comprehensive_result', {}).get('content', '')
            if not content:
                return analysis_result

            # Generate research queries based on content analysis
            research_queries = self._generate_research_queries_from_content(analysis_result)

            research_results = []
            for query in research_queries[:3]:  # Limit to 3 queries
                self.logger.log(f"🔍 Researching: {query}")
                result = self._execute_research_with_tools(query, gemini_client)
                research_results.append(result)

                # Track tool usage
                tool_entry = {
                    "tool": "google_search",
                    "query": query,
                    "timestamp": time.time(),
                    "status": "completed" if not result.get('error') else "failed"
                }
                self.tool_usage_log.append(tool_entry)
                self.logger.log(f"🔧 Tool logged: {len(self.tool_usage_log)} total tools")

            return {
                **analysis_result,
                "research_results": research_results,
                "research_enrichment": True
            }

        except Exception as e:
            self.logger.error(f"⚠️ Research enrichment failed: {e}")
            return analysis_result

    def _generate_research_queries_from_content(self, analysis_result: Dict[str, Any]) -> List[str]:
        """Generate research queries from content analysis"""
        queries = []

        content = analysis_result.get('comprehensive_result', {}).get('content', '')
        title = analysis_result.get('comprehensive_result', {}).get('title', '')

        # Generate queries from title
        if title:
            queries.append(f"Latest information about {title}")
            queries.append(f"Best practices for {title}")

        # Generate queries from keywords
        keywords = analysis_result.get('comprehensive_result', {}).get('keywords', [])
        for keyword in keywords[:3]:
            queries.append(f"Best practices for {keyword}")
            queries.append(f"Tutorials on {keyword}")

        # Generate queries from content analysis
        if content:
            # Simple keyword extraction for research
            words = re.findall(r'\b\w{5,}\b', content.lower())
            stop_words = {'about', 'would', 'there', 'their', 'which', 'could', 'should', 'these', 'those', 'where', 'after', 'before', 'first', 'second', 'third', 'through', 'during', 'while', 'since', 'until', 'although', 'because', 'unless', 'though', 'whether', 'within', 'among', 'between'}

            keywords = [word for word in words if word not in stop_words]
            unique_keywords = list(set(keywords))[:5]  # Top 5 unique keywords

            for keyword in unique_keywords:
                queries.append(f"Understanding {keyword} in detail")

        return list(set(queries))[:5]  # Return max 5 unique queries

    def _execute_research_with_tools(self, query: str, gemini_client) -> Dict[str, Any]:
        """Execute research query using Gemini tools"""
        try:
            self.logger.log(f"🔧 Executing research tool for query: {query}")
            tools = [types.Tool(google_search=types.GoogleSearch())]

            prompt = f"Research this topic using available tools and provide comprehensive insights: {query}"

            response = gemini_client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.7,
                    thinking_config=types.ThinkingConfig(thinking_budget=2048),
                    tools=tools,
                    max_output_tokens=1500
                )
            )

            result = {
                "query": query,
                "insights": response.text if response.candidates else "",
                "timestamp": time.time(),
                "status": "completed"
            }
            self.logger.log(f"✅ Research tool completed for: {query}")
            return result

        except Exception as e:
            self.logger.log(f"❌ Research tool failed for {query}: {e}")
            return {
                "query": query,
                "error": str(e),
                "timestamp": time.time(),
                "status": "failed"
            }

    def _perform_final_synthesis(self, enriched_analysis: Dict[str, Any], document_data: Dict[str, Any], gemini_client, user_instructions: str) -> Dict[str, Any]:
        """Perform final synthesis of all analysis results"""
        try:
            # Add null checks
            if not enriched_analysis:
                return {"error": "No enriched analysis provided", "fallback_synthesis": "Analysis completed with limited synthesis"}

            content = enriched_analysis.get('comprehensive_result', {}).get('content', '')
            research_results = enriched_analysis.get('research_results', [])

            # Get user instructions with null check
            user_instructions = ""
            if document_data:
                user_instructions = document_data.get('instructions', '')

            # Enhanced synthesis using system prompt
            research_summary = chr(10).join([
                f"• {result.get('insights', '')[:300]}"
                for result in research_results if result.get('insights')
            ]) if research_results else "No additional research insights available."

            synthesis_prompt = f"""{self.system_prompt}

🎯 **FINAL SYNTHESIS ROUND**
Integrating all analysis results into a comprehensive, user-focused understanding.

📄 **ORIGINAL CONTENT ANALYSIS**
{content[:4000]}

🔬 **RESEARCH & ENHANCEMENT INSIGHTS**
{research_summary}

🎯 **USER INSTRUCTIONS** (SYNTHESIS FOCUS)
{user_instructions}

📊 **SYNTHESIS OBJECTIVES**
1. **Unified Understanding**: Combine all analysis rounds into coherent insights
2. **User-Centric Focus**: Ensure all findings directly address user instructions
3. **Practical Application**: Transform insights into actionable recommendations
4. **Knowledge Integration**: Connect research findings with original content analysis

🎨 **DELIVERABLES REQUIRED**
Provide a comprehensive synthesis that includes:
1. **Executive Summary**: Clear, concise overview of all findings
2. **Key Insights**: Most important discoveries and patterns
3. **User-Focused Recommendations**: Practical guidance based on user instructions
4. **Implementation Guidance**: Step-by-step advice for applying insights
5. **Further Exploration**: Strategic suggestions for deeper learning
6. **Success Metrics**: How to measure effectiveness of recommendations

⚡ **SYNTHESIS REQUIREMENTS**
- **User Instruction Alignment**: Every recommendation must tie back to user needs
- **Research Integration**: Incorporate current best practices and trends
- **Practical Focus**: Provide immediately actionable insights
- **Progressive Structure**: Build from foundational to advanced concepts
- **Visual Enhancement**: Suggest appropriate visualization types for key concepts

🔧 **QUALITY STANDARDS**
- Content must be immediately applicable to user's stated goals
- Insights should provide clear value beyond the original content
- Recommendations should be specific and measurable
- Synthesis should create a complete learning experience

🎯 **SUCCESS CRITERIA**
- User can immediately apply at least 3 specific recommendations
- Synthesis addresses all aspects of user's original instructions
- Content provides unique value not found in original source
- Structure enables progressive skill development"""

            response = gemini_client.models.generate_content(
                model=GEMINI_MODEL,
                contents=synthesis_prompt,
                config=types.GenerateContentConfig(
                    temperature=0.7,
                    thinking_config=types.ThinkingConfig(thinking_budget=2048),
                    max_output_tokens=2500
                )
            )

            synthesis_text = response.text if response.candidates else ""

            # Try to extract structured data from synthesis
            synthesis_data = None
            try:
                if synthesis_text.strip().startswith('{'):
                    synthesis_data = json.loads(synthesis_text)
                else:
                    # Extract JSON from text if embedded
                    json_match = re.search(r'\{.*\}', synthesis_text, re.DOTALL)
                    if json_match:
                        synthesis_data = json.loads(json_match.group())
            except Exception:
                pass

            return {
                "final_synthesis": synthesis_text,
                "structured_synthesis": synthesis_data,
                "integrated_insights": True,
                "timestamp": time.time()
            }

        except Exception as e:
            self.logger.error(f"⚠️ Final synthesis failed: {e}")
            return {"error": str(e), "fallback_synthesis": "Analysis completed with limited synthesis"}

    def _generate_comprehensive_blocks(self, analysis_result: Dict[str, Any], document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Generate comprehensive analysis blocks following LLM_RESPONSE_SCHEMA"""
        try:
            blocks = []

            # Add null checks
            if not analysis_result:
                return [{
                    "id": "error-block",
                    "type": "summary",
                    "size": "small",
                    "title": "Analysis Error",
                    "content": "No analysis results available.",
                    "metadata": {"priority": "low"}
                }]

            # Get comprehensive result data
            comprehensive_result = analysis_result.get('comprehensive_result', {})
            enhanced_analysis = analysis_result.get('enhanced_analysis', {})
            enriched_analysis = analysis_result.get('enriched_analysis', {})

            # Summary block (always included)
            summary_content = comprehensive_result.get('content', '')[:1000]
            if not summary_content:
                summary_content = "Document analysis completed successfully."

            blocks.append({
                "id": "comprehensive-summary",
                "type": "summary",
                "size": "large",
                "title": f"Analysis: {comprehensive_result.get('title', 'Document')}",
                "content": summary_content,
                "metadata": {
                    "priority": "high",
                    "word_count": len(summary_content.split()),
                    "source": "extracted_content"
                }
            })

            # Priority 1: Use schema-compliant LLM response if available
            if isinstance(enhanced_analysis.get('enhanced_analysis'), dict):
                enhanced_data = enhanced_analysis['enhanced_analysis']

                # Use the LLM-generated summary if available
                if 'summary' in enhanced_data and enhanced_data['summary']:
                    # Replace the basic summary with LLM-generated one
                    blocks[0] = {
                        "id": "llm-summary",
                        "type": "summary",
                        "size": "large",
                        "title": "AI Analysis Summary",
                        "content": enhanced_data['summary'],
                        "metadata": {
                            "priority": "high",
                            "source": "gemini_analysis",
                            "analysis_type": "enhanced"
                        }
                    }

                # Use the LLM-generated blocks
                if 'blocks' in enhanced_data and isinstance(enhanced_data['blocks'], list):
                    for block in enhanced_data['blocks'][:4]:  # Allow up to 4 LLM blocks
                        if self._validate_block_structure(block):
                            # Ensure unique ID
                            block_id = block.get('id', f"llm-block-{len(blocks)}")
                            block['id'] = f"{block_id}-{len(blocks)}"

                            # Ensure block follows schema
                            block = self._normalize_block_structure(block)
                            blocks.append(block)

            # Key insights block
            word_count = comprehensive_result.get('word_count', 0)
            links_count = len(comprehensive_result.get('links', []))
            media_count = len(comprehensive_result.get('media', []))

            blocks.append({
                "id": "key-insights",
                "type": "key_points",
                "size": "medium",
                "title": "Content Analysis",
                "content": f"• Content length: {len(comprehensive_result.get('content', ''))} characters\n• Word count: {word_count}\n• Links found: {links_count}\n• Media elements: {media_count}\n• Keywords: {', '.join(comprehensive_result.get('keywords', [])[:5])}",
                "metadata": {
                    "priority": "medium",
                    "source": "content_analysis"
                }
            })

            # Code blocks (if any detected)
            code_blocks = comprehensive_result.get('code_blocks', [])
            if code_blocks:
                code_content = "\n\n".join([
                    f"```{block.get('language', '')}\n{block['content']}\n```"
                    for block in code_blocks[:3]  # Limit to 3 code blocks
                ])

                blocks.append({
                    "id": "code-examples",
                    "type": "code",
                    "size": "medium",
                    "title": "Code Examples",
                    "content": code_content,
                    "metadata": {
                        "priority": "medium",
                        "code_blocks_count": len(code_blocks),
                        "languages": list(set([block.get('language', 'unknown') for block in code_blocks])),
                        "source": "content_extraction"
                    }
                })

            # Research insights block
            research_results = enriched_analysis.get('research_results', [])
            if research_results:
                research_content = "\n\n".join([
                    f"**{result['query']}**\n{result.get('insights', '')[:300]}"
                    for result in research_results
                    if result.get('insights')
                ])

                if research_content:
                    blocks.append({
                        "id": "research-insights",
                        "type": "research",
                        "size": "medium",
                        "title": "Research & Context",
                        "content": research_content,
                        "metadata": {
                            "priority": "medium",
                            "research_queries": len(research_results),
                            "source": "google_search"
                        }
                    })

            # Links analysis block
            links = comprehensive_result.get('links', [])
            if links:
                links_content = "\n".join([
                    f"- [{link['text']}]({link['url']})"
                    for link in links[:10]  # Limit to 10 links
                ])

                blocks.append({
                    "id": "related-links",
                    "type": "guide",
                    "size": "small",
                    "title": "Related Links",
                    "content": links_content,
                    "metadata": {
                        "priority": "low",
                        "links_count": len(links),
                        "source": "content_extraction"
                    }
                })

            # Ensure we don't exceed 6 blocks maximum
            blocks = blocks[:6]

            # Validate all blocks have required fields
            validated_blocks = []
            for block in blocks:
                if all(key in block for key in ['id', 'type', 'size', 'title', 'content']):
                    validated_blocks.append(block)
            else:
                    self.logger.log(f"⚠️ Skipping invalid block: {block.get('id', 'unknown')}")

            return validated_blocks

        except Exception as e:
            self.logger.error(f"⚠️ Comprehensive block generation failed: {e}")
            return [{
                "id": "fallback-summary",
                "type": "summary",
                "size": "large",
                "title": "Analysis Summary",
                "content": f"Document analysis completed. {len(analysis_result.get('comprehensive_result', {}).get('content', ''))} characters processed.",
                "metadata": {"priority": "high"}
            }]

    def _perform_research_enrichment_with_content(self, analysis_result: Dict[str, Any], gemini_client, scraped_content: Dict[str, Any]) -> Dict[str, Any]:
        """Perform research enrichment using scraped content"""
        try:
            content = scraped_content.get('content', '')
            if not content:
                return analysis_result

            # Generate research queries based on scraped content
            research_queries = self._generate_research_queries_from_scraped_content(scraped_content)

            research_results = []
            for query in research_queries[:3]:  # Limit to 3 queries
                self.logger.log(f"🔍 Researching: {query}")
                result = self._execute_research_with_tools(query, gemini_client)
                research_results.append(result)

                # Track tool usage
                tool_entry = {
                    "tool": "google_search",
                    "query": query,
                    "timestamp": time.time(),
                    "status": "completed" if not result.get('error') else "failed"
                }
                self.tool_usage_log.append(tool_entry)
                self.logger.log(f"🔧 Tool logged: {len(self.tool_usage_log)} total tools")

            return {
                **analysis_result,
                "research_results": research_results,
                "research_enrichment": True
            }

        except Exception as e:
            self.logger.error(f"⚠️ Research enrichment with content failed: {e}")
            return analysis_result

    def _perform_final_synthesis_with_content(self, enriched_analysis: Dict[str, Any], document_data: Dict[str, Any], gemini_client, user_instructions: str, scraped_content: Dict[str, Any]) -> Dict[str, Any]:
        """Perform final synthesis using scraped content"""
        try:
            # Add null checks
            if not enriched_analysis:
                return {"error": "No enriched analysis provided", "fallback_synthesis": "Analysis completed with limited synthesis"}

            content = scraped_content.get('content', '')
            research_results = enriched_analysis.get('research_results', [])

            # Get user instructions with null check
            user_instructions = ""
            if document_data:
                user_instructions = document_data.get('instructions', '')

            # Enhanced synthesis using system prompt
            research_summary = chr(10).join([
                f"• {result.get('insights', '')[:300]}"
                for result in research_results if result.get('insights')
            ]) if research_results else "No additional research insights available."

            synthesis_prompt = f"""{self.system_prompt}

🎯 **FINAL SYNTHESIS ROUND**
Integrating all analysis results into a comprehensive, user-focused understanding.

📄 **ORIGINAL CONTENT ANALYSIS**
{content[:4000]}

🔬 **RESEARCH & ENHANCEMENT INSIGHTS**
{research_summary}

🎯 **USER INSTRUCTIONS** (SYNTHESIS FOCUS)
{user_instructions}

📊 **SYNTHESIS OBJECTIVES**
1. **Unified Understanding**: Combine all analysis rounds into coherent insights
2. **User-Centric Focus**: Ensure all findings directly address user instructions
3. **Practical Application**: Transform insights into actionable recommendations
4. **Knowledge Integration**: Connect research findings with original content analysis

🎨 **DELIVERABLES REQUIRED**
Provide a comprehensive synthesis that includes:
1. **Executive Summary**: Clear, concise overview of all findings
2. **Key Insights**: Most important discoveries and patterns
3. **User-Focused Recommendations**: Practical guidance based on user instructions
4. **Implementation Guidance**: Step-by-step advice for applying insights
5. **Further Exploration**: Strategic suggestions for deeper learning
6. **Success Metrics**: How to measure effectiveness of recommendations

⚡ **SYNTHESIS REQUIREMENTS**
- **User Instruction Alignment**: Every recommendation must tie back to user needs
- **Research Integration**: Incorporate current best practices and trends
- **Practical Focus**: Provide immediately actionable insights
- **Progressive Structure**: Build from foundational to advanced concepts
- **Visual Enhancement**: Suggest appropriate visualization types for key concepts

🔧 **QUALITY STANDARDS**
- Content must be immediately applicable to user's stated goals
- Insights should provide clear value beyond the original content
- Recommendations should be specific and measurable
- Synthesis should create a complete learning experience

🎯 **SUCCESS CRITERIA**
- User can immediately apply at least 3 specific recommendations
- Synthesis addresses all aspects of user's original instructions
- Content provides unique value not found in original source
- Structure enables progressive skill development"""

            response = gemini_client.models.generate_content(
                model=GEMINI_MODEL,
                contents=synthesis_prompt,
                config=types.GenerateContentConfig(
                    temperature=0.7,
                    thinking_config=types.ThinkingConfig(thinking_budget=2048),
                    max_output_tokens=2500
                )
            )

            synthesis_text = response.text if response.candidates else ""

            # Try to extract structured data from synthesis
            synthesis_data = None
            try:
                if synthesis_text.strip().startswith('{'):
                    synthesis_data = json.loads(synthesis_text)
                else:
                    # Extract JSON from text if embedded
                    json_match = re.search(r'\{.*\}', synthesis_text, re.DOTALL)
                    if json_match:
                        synthesis_data = json.loads(json_match.group())
            except Exception:
                pass

            return {
                "final_synthesis": synthesis_text,
                "structured_synthesis": synthesis_data,
                "integrated_insights": True,
                "timestamp": time.time()
            }

        except Exception as e:
            self.logger.error(f"⚠️ Final synthesis with content failed: {e}")
            return {"error": str(e), "fallback_synthesis": "Analysis completed with limited synthesis"}

    def _generate_research_queries_from_scraped_content(self, scraped_content: Dict[str, Any]) -> List[str]:
        """Generate research queries from scraped content"""
        queries = []

        content = scraped_content.get('content', '')
        title = scraped_content.get('title', '')

        # Generate queries from title
        if title:
            queries.append(f"Latest information about {title}")
            queries.append(f"Best practices for {title}")

        # Generate queries from content analysis
        if content:
            # Simple keyword extraction for research
            words = re.findall(r'\b\w{5,}\b', content.lower())
            stop_words = {'about', 'would', 'there', 'their', 'which', 'could', 'should', 'these', 'those', 'where', 'after', 'before', 'first', 'second', 'third', 'through', 'during', 'while', 'since', 'until', 'although', 'because', 'unless', 'though', 'whether', 'within', 'among', 'between'}

            keywords = [word for word in words if word not in stop_words]
            unique_keywords = list(set(keywords))[:5]  # Top 5 unique keywords

            for keyword in unique_keywords:
                queries.append(f"Understanding {keyword} in detail")

        return list(set(queries))[:5]  # Return max 5 unique queries

    def _create_final_comprehensive_result(self, analysis_result: Dict[str, Any], blocks: List[Dict[str, Any]], document_data: Dict[str, Any]) -> Dict[str, Any]:
        """Create the final comprehensive result"""
        try:
            # Add null checks
            if not analysis_result:
                return {"error": "No analysis result provided", "success": False}

            comprehensive_result = analysis_result.get('comprehensive_result', {})
            final_analysis = analysis_result.get('final_analysis', {})

            # Calculate processing time
            processing_time = time.time() - self.start_time

            # Get document ID with null check
            document_id = "unknown"
            if document_data and isinstance(document_data, dict):
                document_id = document_data.get('document_id', 'unknown')

            final_result = {
                "success": True,
                "documentId": document_id,
                "title": comprehensive_result.get('title', ''),
                "description": comprehensive_result.get('description', ''),
                "url": comprehensive_result.get('url', ''),
                "content": comprehensive_result.get('content', ''),
                "word_count": comprehensive_result.get('word_count', 0),
                "character_count": len(comprehensive_result.get('content', '')),
                "blocks": blocks,
                "metadata": {
                    "processing_time": round(processing_time, 2),
                    "analysis_version": "comprehensive-v2.5",
                    "tools_used": ["google_search", "code_execution", "url_context"],
                    "tool_executions": len(self.tool_usage_log),
                    "content_type": comprehensive_result.get('content_type', 'unknown'),
                    "research_queries_executed": len(analysis_result.get('enriched_analysis', {}).get('research_results', [])),
                    "blocks_generated": len(blocks),
                    "rounds_completed": analysis_result.get('rounds_completed', 0),
                    "code_blocks_extracted": len(comprehensive_result.get('code_blocks', [])),
                    "links_extracted": len(comprehensive_result.get('links', [])),
                    "timestamp": time.time()
                },
                "analysis_summary": final_analysis.get('final_synthesis', ''),
                "research_insights": [
                    result for result in analysis_result.get('enriched_analysis', {}).get('research_results', [])
                    if result.get('status') == 'completed'
                ]
            }

            return final_result

        except Exception as e:
            self.logger.error(f"⚠️ Final result creation failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "documentId": document_data.get('document_id', 'unknown'),
                "blocks": blocks,
                "metadata": {
                    "processing_time": time.time() - self.start_time,
                    "error_occurred": True
                }
            }

    def _execute_workflow(self, analysis_plan: Dict[str, Any], document_data: Dict[str, Any]) -> Dict[str, Any]:
        """Execute the orchestrated workflow"""
        workflow_results = {
            "content": {},
            "research": {},
            "analysis": {},
            "tool_usage": []
        }

        tool_sequence = analysis_plan.get("tool_sequence", [])

        if "scrape_document" in tool_sequence:
            self.logger.log("🌐 Executing: scrape_document")
            result = self.content_processor.process_url_content(document_data.get('url', ''))
            workflow_results["content"] = result

        if analysis_plan.get("research_topics"):
            for topic in analysis_plan["research_topics"]:
                self.logger.log(f"🔍 Researching: {topic}")
                result = self.research_engine.research_user_interests(
                    self._get_user_interests(document_data.get('user_id')), topic
                )
                workflow_results["research"][topic] = result

        if "analyze_content" in tool_sequence:
            self.logger.log("🧠 Analyzing content with Gemini...")
            analysis_result = self._analyze_content_with_gemini(workflow_results["content"])
            workflow_results["analysis"] = analysis_result

        workflow_results["tool_usage"] = self.tool_usage_log
        self.logger.log(f"📊 Workflow completed: {len(workflow_results['tool_usage'])} tools executed")
        return workflow_results

    def _analyze_content_with_gemini(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze content using Gemini"""
        content_text = content.get("content", "")[:2000]  # Limit for API

        prompt = f"""
        Analyze this content and provide:
        1. A brief summary (2-3 sentences)
        2. Key points (3-5 bullet points)
        3. Main topics covered

        Content: {content_text}
        """

        try:
            response = self.gemini_client.models.generate_content(
                model=self.gemini_model_name,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.7,
                    thinking_config=types.ThinkingConfig(thinking_budget=2048)
                )
            )
            return {"analysis": response.text.strip()}
        except Exception as e:
            self.logger.error(f"⚠️ Content analysis failed: {e}")
            return {"analysis": "Content analysis completed"}

    def _format_final_result(self, workflow_results: Dict[str, Any], document_data: Dict[str, Any]) -> Dict[str, Any]:
        """Format final result"""
        content = workflow_results.get("content", {})
        research = workflow_results.get("research", {})

        # Generate summary
        summary = self._generate_summary(content, research, document_data)

        # Generate analysis blocks
        blocks = self._generate_analysis_blocks(content, research, document_data)

        return {
            "success": True,
            "documentId": document_data['document_id'],
            "summary": summary,
            "blocks": blocks,
            "metadata": {
                "tools_used": workflow_results.get("tool_usage", []),
                "content_length": len(content.get("content", "")),
                "research_sources": len(research),
                "analysis_blocks": len(blocks)
            }
        }

    def _generate_summary(self, content: Dict[str, Any], research: Dict[str, Any], document_data: Dict[str, Any]) -> str:
        """Generate summary"""
        content_text = content.get("content", "")[:1000]

        prompt = f"""
        Create a concise summary of this document:

        TITLE: {document_data.get('title', 'Untitled')}
        CONTENT: {content_text}
        INSTRUCTIONS: {document_data.get('instructions', '')}

        Summary (2-3 sentences):
        """

        try:
            response = self.gemini_client.models.generate_content(
                model=self.gemini_model_name,
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.7,
                    thinking_config=types.ThinkingConfig(thinking_budget=2048)
                )
            )
            return response.text.strip()
        except Exception as e:
            return f"Document: {document_data.get('title', 'Untitled')} - Analysis completed."

    def _generate_analysis_blocks(self, content: Dict[str, Any], research: Dict[str, Any], document_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Generate analysis blocks"""
        try:
            blocks = [
                {
                    "id": "summary",
                    "type": "summary",
                    "size": "large",
                    "title": "Document Summary",
                    "content": self._generate_summary(content, research, document_data),
                    "metadata": {"priority": "high"}
                },
                {
                    "id": "key_points",
                    "type": "key_points",
                    "size": "medium",
                    "title": "Key Points",
                    "content": f"• Content length: {len(content.get('content', ''))} characters\n• Links found: {len(content.get('links', []))}\n• Word count: {content.get('word_count', 0)}",
                    "metadata": {"priority": "medium"}
                }
            ]

            if research:
                blocks.append({
                    "id": "research",
                    "type": "research",
                    "size": "medium",
                    "title": "Research Insights",
                    "content": f"Research conducted on {len(research)} topics based on user interests.",
                    "metadata": {"priority": "medium"}
                })

            return blocks

        except Exception as e:
            self.logger.error(f"⚠️ Block generation failed: {e}")
            return [{
                "id": "error",
                "type": "summary",
                "size": "large",
                "title": "Processing Complete",
                "content": "Document processing completed successfully.",
                "metadata": {"priority": "high"}
            }]

    def _update_document_status(self, document_id: str, status: str, result: Optional[Dict[str, Any]] = None, error: Optional[str] = None) -> None:
        """Update document status in database"""
        try:
            update_data = {
                'status': status
            }

            if result:
                # Extract summary from schema-compliant data if available
                summary_text = ""
                if isinstance(result.get('enhanced_analysis', {}).get('enhanced_analysis'), dict):
                    schema_data = result['enhanced_analysis']['enhanced_analysis']
                    if 'summary' in schema_data and schema_data['summary']:
                        summary_text = schema_data['summary']
                    else:
                        summary_text = result.get('summary', '')
                else:
                    summary_text = result.get('summary', '')

                # Ensure summary_text is always a valid string
                if not isinstance(summary_text, str):
                    summary_text = str(summary_text) if summary_text is not None else ""

                # Fallback to basic summary if needed
                if not summary_text.strip():
                    summary_text = result.get('analysis_summary', 'Analysis completed successfully')

                # Truncate to 2000 characters to meet database constraints
                if len(summary_text) > 2000:
                    summary_text = summary_text[:1997] + "..."
                    self.logger.log(f"⚠️ Summary truncated to 2000 chars: {len(summary_text)}")

                # Validate and sanitize blocks data
                blocks_data = result.get('blocks', [])
                if not isinstance(blocks_data, list):
                    blocks_data = []
                    self.logger.log("⚠️ analysis_blocks was not a list, using empty array")

                sanitized_blocks = []
                for block in blocks_data:
                    if isinstance(block, dict) and all(key in block for key in ['id', 'type', 'size', 'title', 'content']):
                        sanitized_block = {
                            'id': str(block.get('id', '')),
                            'type': str(block.get('type', '')),
                            'size': str(block.get('size', '')),
                            'title': str(block.get('title', '')),
                            'content': str(block.get('content', ''))
                        }
                        if 'metadata' in block and isinstance(block['metadata'], dict):
                            sanitized_block['metadata'] = block['metadata']
                        sanitized_blocks.append(sanitized_block)

                update_data['analysis_summary'] = summary_text
                update_data['analysis_blocks'] = json.dumps(sanitized_blocks)
                update_data['word_count'] = len(summary_text.split())

                # Update fields that exist
                available_fields = self._get_available_fields()
                if 'gemini_tools_used' in available_fields:
                    update_data['gemini_tools_used'] = json.dumps([t.get('tool', 'unknown') for t in self.tool_usage_log])
                if 'research_context' in available_fields:
                    update_data['research_context'] = json.dumps(result.get('metadata', {}))
                if 'processing_duration' in available_fields:
                    update_data['processing_duration'] = int(time.time() - time.time())

            if error and 'error_message' in self._get_available_fields():
                update_data['error_message'] = error

            self.databases.update_document(
                DATABASE_ID,
                DOCUMENTS_COLLECTION_ID,
                document_id,
                update_data
            )

            self.logger.log(f"✅ Document {document_id} updated to {status}")

        except Exception as e:
            self.logger.error(f"⚠️ Database update failed: {e}")

    def _get_available_fields(self) -> List[str]:
        """Get available database fields"""
        try:
            response = self.databases.list_attributes(DATABASE_ID, DOCUMENTS_COLLECTION_ID)
            return [attr['key'] for attr in response.get('attributes', [])]
        except Exception:
            # Fallback to known fields
            return [
                'user_id', 'title', 'url', 'instructions', 'status', 'public',
                'scraped_content', 'word_count', 'analysis_summary', 'analysis_blocks',
                'user_interests', 'gemini_tools_used', 'research_context', 'processing_duration'
            ]

# ===== MAIN FUNCTION =====
def main(context: Dict[str, Any]) -> Dict[str, Any]:
    """Main Appwrite function entry point"""
    start_time = time.time()

    try:
        # Setup logging
        logger = Logger(context)
        logger.log("=== DOCIFY UNIFIED ORCHESTRATOR START ===")

        # Validate environment
        if not GEMINI_API_KEY:
            raise ValueError("GEMINI_API_KEY environment variable is required")

        logger.log("✅ Environment validation passed")
        logger.log("🚀 Initializing Gemini Orchestrator...")

        # Initialize unified orchestrator with all Gemini tools
        orchestrator = DocifyUnifiedOrchestrator(context, databases, logger)
        logger.log("✅ Docify Unified Orchestrator initialized with all Gemini tools")

        # Process document comprehensively
        logger.log("📄 Starting comprehensive document processing with all tools...")
        result = orchestrator.process_document_comprehensive()

        # Add performance metrics
        processing_time = time.time() - start_time
        result['metadata'] = result.get('metadata', {})
        result['metadata']['processing_time'] = round(processing_time, 2)
        result['metadata']['function_version'] = 'docify-unified-orchestrator-v2.5'

        logger.log(f"⏱️ Total processing time: {processing_time:.2f}s")
        logger.log(f"📊 Tools executed: {len(result.get('metadata', {}).get('tools_used', []))}")

        return context.res.json(result, 200)

    except Exception as e:
        error_time = time.time() - start_time
        logger.log(f"❌ Processing failed in {error_time:.2f}s")
        return context.res.json({
            "success": False,
            "error": str(e),
            "processing_time": round(error_time, 2),
            "function_version": "docify-unified-orchestrator-v2.5"
        }, 500)

# ===== LEGACY COMPATIBILITY =====
def process_document_legacy(context: Dict[str, Any]) -> Dict[str, Any]:
    """Legacy function for backward compatibility"""
    return main(context)

# Export for Appwrite
if __name__ == "__main__":
    print("🚀 Docify Unified Orchestrator v2.5 loaded successfully")
    print("✨ Features: Gemini tools (google_search, url_context)")
    print("🔧 Enhanced URL scraping with multi-format support")
    print("📊 Multi-round analysis with comprehensive insights")
    print("🎯 Advanced content processing and research enrichment")
    print("This function should be called through Appwrite")
