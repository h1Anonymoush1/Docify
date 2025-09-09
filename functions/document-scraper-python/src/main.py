import os
from appwrite.client import Client
from appwrite.services.databases import Databases
from appwrite.id import ID
import re
from urllib.parse import urlparse, urljoin
import json
from collections import defaultdict
import requests
from bs4 import BeautifulSoup
import chardet
import PyPDF2
import docx
import pandas as pd
import feedparser
from io import BytesIO
import base64
import mimetypes
import time

# BROWSERLESS_API_KEY for JavaScript rendering (loaded from environment)
BROWSERLESS_API_KEY = os.environ.get('BROWSERLESS_API_KEY', '')

def smart_fetch_html(url, timeout=15):
    """Fetch HTML content using multiple strategies like url-to-markdown approach."""
    html_content = None
    fetch_method = 'basic'

    # Fetch strategies (adapted from url-to-markdown)
    fetch_strategies = [
        {
            'name': 'modern-browser',
            'headers': {
                'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
                'Accept-Encoding': 'gzip, deflate, br',
                'DNT': '1',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
                'Sec-Fetch-Dest': 'document',
                'Sec-Fetch-Mode': 'navigate',
                'Sec-Fetch-Site': 'none',
                'Cache-Control': 'max-age=0'
            }
        },
        {
            'name': 'mobile',
            'headers': {
                'User-Agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 17_0 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.0 Mobile/15E148 Safari/604.1',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
                'Accept-Encoding': 'gzip, deflate, br'
            }
        },
        {
            'name': 'bot-friendly',
            'headers': {
                'User-Agent': 'Mozilla/5.0 (compatible; DocifyBot/1.0; +https://docify.app)',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9'
            }
        }
    ]

    # Try each strategy
    for strategy in fetch_strategies:
        try:
            print(f"Trying fetch strategy: {strategy['name']}")
            response = requests.get(url, headers=strategy['headers'], timeout=timeout)

            if response.status_code == 200:
                html_content = response.text
                fetch_method = strategy['name']
                print(f"Successfully fetched {len(html_content)} characters using {strategy['name']} strategy")

                # Check if content seems too small (likely JS-dependent)
                if len(html_content) < 10000:
                    print(f"Content seems small ({len(html_content)} chars), might be JS-dependent. Will try rendering services.")
                    continue
                break
            else:
                print(f"{strategy['name']} failed with status: {response.status_code}")
        except Exception as strategy_error:
            print(f"{strategy['name']} failed: {strategy_error}")

    # Try JS rendering services if content is small or all strategies failed
    tried_rendering = False
    if not html_content or len(html_content or '') < 10000:
        tried_rendering = True
        print('Content seems small or no content fetched, trying JS rendering services...')

        # Try Browserless.io first (if API key is available)
        if BROWSERLESS_API_KEY:
            try:
                print('Trying Browserless.io...')
                browserless_url = f"https://production-sfo.browserless.io/content?token={BROWSERLESS_API_KEY}"
                response = requests.post(browserless_url, json={
                    'url': url,
                    'gotoOptions': {
                        'waitUntil': 'networkidle2',
                        'timeout': 30000
                    }
                }, timeout=30)

                if response.status_code == 200:
                    rendered_html = response.text
                    if len(rendered_html) > len(html_content or ''):
                        html_content = rendered_html
                        fetch_method = 'browserless-io'
                        print(f"Successfully rendered {len(html_content)} characters using Browserless.io")
                    else:
                        print(f"Browserless.io returned same or smaller content ({len(rendered_html)} chars)")
                else:
                    print(f"Browserless.io failed with status: {response.status} - {response.text}")
            except Exception as browserless_error:
                print(f"Browserless.io error: {browserless_error}")
        else:
            print("BROWSERLESS_API_KEY not provided, skipping Browserless.io")

        # Fallback to free services if Browserless.io didn't help
        if not html_content or len(html_content) < 10000:
            free_services = [
                {
                    'name': 'rendertron-free',
                    'url': f"https://render-tron.appspot.com/render/{url}",
                    'type': 'direct'
                }
            ]

            for service in free_services:
                try:
                    print(f"Trying fallback {service['name']}...")
                    response = requests.get(service['url'], headers={
                        'User-Agent': 'Mozilla/5.0 (compatible; DocifyBot/1.0)'
                    }, timeout=20)

                    if response.status_code == 200:
                        rendered_html = response.text
                        if len(rendered_html) > len(html_content or ''):
                            html_content = rendered_html
                            fetch_method = service['name']
                            print(f"Successfully rendered {len(html_content)} characters using {service['name']}")
                            break
                        else:
                            print(f"{service['name']} returned same or smaller content ({len(rendered_html)} chars)")
                except Exception as service_error:
                    print(f"{service['name']} failed: {service_error}")

    return html_content, fetch_method, tried_rendering


def detect_content_type_from_response(response, url):
    """Detect the content type of a response."""
    # Check Content-Type header
    content_type_header = response.headers.get('Content-Type', '').lower()

    # Check URL extension
    url_path = urlparse(url).path.lower()

    # Determine content type
    if 'application/pdf' in content_type_header or url_path.endswith('.pdf'):
        return 'pdf'
    elif 'application/msword' in content_type_header or url_path.endswith(('.doc', '.docx')):
        return 'doc'
    elif 'application/vnd.ms-excel' in content_type_header or url_path.endswith(('.xls', '.xlsx')):
        return 'excel'
    elif 'text/csv' in content_type_header or url_path.endswith('.csv'):
        return 'csv'
    elif 'application/json' in content_type_header or url_path.endswith('.json'):
        return 'json'
    elif 'application/xml' in content_type_header or url_path.endswith(('.xml', '.rss', '.atom')):
        return 'xml'
    elif 'text/plain' in content_type_header or url_path.endswith(('.txt', '.md')):
        return 'text'
    else:
        # Default to HTML for web pages
        return 'html'

def extract_content_by_type(response, content_type, url):
    """Extract content based on the detected content type."""
    try:
        if content_type == 'pdf':
            return extract_pdf_content(response, url)
        elif content_type == 'doc':
            return extract_doc_content(response, url)
        elif content_type == 'excel':
            return extract_excel_content(response, url)
        elif content_type == 'csv':
            return extract_csv_content(response, url)
        elif content_type == 'json':
            return extract_json_content(response, url)
        elif content_type == 'xml':
            return extract_xml_content(response, url)
        elif content_type == 'text':
            return extract_text_content(response, url)
        else:  # Default to HTML
            return extract_html_content(response, url)
    except Exception as e:
        print(f"Error extracting {content_type} content from {url}: {e}")
        return None

def extract_pdf_content(response, url):
    """Extract content from PDF files."""
    try:
        pdf_file = BytesIO(response.content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        content = []
        for page in pdf_reader.pages[:10]:  # Limit to first 10 pages
            text = page.extract_text()
            if text.strip():
                content.append(text)

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'PDF Document - {len(pdf_reader.pages)} pages',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'pdf',
            'metadata': {
                'pages': len(pdf_reader.pages),
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting PDF content: {e}")
        return None

def extract_doc_content(response, url):
    """Extract content from Word documents."""
    try:
        doc_file = BytesIO(response.content)
        doc = docx.Document(doc_file)

        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'Word Document - {len(doc.paragraphs)} paragraphs',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'doc',
            'metadata': {
                'paragraphs': len(doc.paragraphs),
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting DOC content: {e}")
        return None

def extract_excel_content(response, url):
    """Extract content from Excel files."""
    try:
        excel_file = BytesIO(response.content)
        df = pd.read_excel(excel_file)

        # Convert DataFrame to readable text
        content = []
        content.append(f"Sheet: {df.columns.tolist()}")
        content.append("Data Preview:")
        content.append(str(df.head(20)))  # First 20 rows

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'Excel Spreadsheet - {len(df)} rows, {len(df.columns)} columns',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'excel',
            'metadata': {
                'rows': len(df),
                'columns': len(df.columns),
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting Excel content: {e}")
        return None

def extract_csv_content(response, url):
    """Extract content from CSV files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        csv_text = response.content.decode(encoding, errors='ignore')
        lines = csv_text.split('\n')[:50]  # First 50 lines

        content = []
        content.append("CSV Data Preview:")
        content.extend(lines[:20])  # First 20 lines

        full_content = '\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'CSV File - {len(lines)} lines',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'csv',
            'metadata': {
                'lines': len(lines),
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting CSV content: {e}")
        return None

def extract_json_content(response, url):
    """Extract content from JSON files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        json_text = response.content.decode(encoding, errors='ignore')
        json_data = json.loads(json_text)

        # Convert JSON to readable text
        content = []
        content.append("JSON Structure:")
        content.append(json.dumps(json_data, indent=2)[:2000])  # First 2000 chars

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'JSON Data - {len(json_text)} characters',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'json',
            'metadata': {
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting JSON content: {e}")
        return None

def extract_xml_content(response, url):
    """Extract content from XML/RSS files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        xml_text = response.content.decode(encoding, errors='ignore')

        # Check if it's an RSS/Atom feed
        if '<rss' in xml_text.lower() or '<feed' in xml_text.lower():
            return extract_feed_content(xml_text, url)

        # Regular XML parsing
        from xml.etree import ElementTree as ET
        root = ET.fromstring(xml_text)

        content = []
        content.append("XML Structure:")
        content.append(xml_to_text(root, level=0)[:2000])

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'XML Document - {len(xml_text)} characters',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'xml',
            'metadata': {
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting XML content: {e}")
        return None

def extract_feed_content(xml_text, url):
    """Extract content from RSS/Atom feeds."""
    try:
        feed = feedparser.parse(xml_text)

        content = []
        content.append(f"Feed Title: {feed.feed.get('title', 'Unknown')}")
        content.append(f"Feed Description: {feed.feed.get('description', '')}")
        content.append(f"Total Entries: {len(feed.entries)}")

        # Add recent entries
        for i, entry in enumerate(feed.entries[:10]):
            content.append(f"\n--- Entry {i+1} ---")
            content.append(f"Title: {entry.get('title', '')}")
            content.append(f"Link: {entry.get('link', '')}")
            if 'summary' in entry:
                content.append(f"Summary: {entry.summary[:500]}...")

        full_content = '\n'.join(content)

        return {
            'url': url,
            'title': feed.feed.get('title', extract_title_from_url(url)),
            'description': f'RSS/Atom Feed - {len(feed.entries)} entries',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'feed',
            'metadata': {
                'entries': len(feed.entries),
                'feed_type': 'rss' if '<rss' in xml_text.lower() else 'atom'
            }
        }
    except Exception as e:
        print(f"Error extracting feed content: {e}")
        return None

def extract_text_content(response, url):
    """Extract content from plain text files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        text_content = response.content.decode(encoding, errors='ignore')

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'Text File - {len(text_content)} characters',
            'content': text_content[:10000],  # Limit to 10K chars
            'word_count': len(text_content.split()),
            'content_type': 'text',
            'metadata': {
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting text content: {e}")
        return None

def extract_html_content(response, url):
    """Extract content from HTML pages."""
    title = extract_title(response, url)
    description = extract_description(response)
    content = extract_main_content(response)

    return {
        'url': url,
        'title': title,
        'description': description,
        'content': content,
        'word_count': len(content.split()) if content else 0,
        'content_type': 'html'
    }

def extract_title_from_url(url):
    """Extract a title from URL when no other title is available."""
    path = urlparse(url).path
    filename = path.split('/')[-1]
    if '.' in filename:
        filename = filename.split('.')[0]
    return filename.replace('-', ' ').replace('_', ' ').title() or 'Document'

def xml_to_text(element, level=0):
    """Convert XML element to readable text."""
    indent = "  " * level
    text = f"{indent}<{element.tag}>"

    if element.text and element.text.strip():
        text += f" {element.text.strip()}"

    for child in element:
        text += "\n" + xml_to_text(child, level + 1)

    if not element.text or not element.text.strip():
        text += f"{indent}</{element.tag}>"

    return text

def extract_title(response, url):
    """Extract page title."""
    title_selectors = [
        'title::text',
        'h1::text',
        'h1 a::text',
        '[property="og:title"]::attr(content)',
        'meta[name="title"]::attr(content)'
    ]

    for selector in title_selectors:
        title = response.css(selector).get()
        if title:
            return title.strip()

    return url.split('/')[-1] or 'Untitled Page'

def extract_description(response):
    """Extract page description."""
    desc_selectors = [
        'meta[name="description"]::attr(content)',
        'meta[property="og:description"]::attr(content)',
        'meta[name="twitter:description"]::attr(content)'
    ]

    for selector in desc_selectors:
        desc = response.css(selector).get()
        if desc:
            return desc.strip()

    return ''

def extract_main_content(response):
    """Extract main content from the page."""
    # Try various content selectors in order of preference
    content_selectors = [
        'main',
        '[role="main"]',
        '.content',
        '.main-content',
        '#content',
        '#main',
        'article',
        '.article-content',
        '.post-content',
        '.entry-content',
        '.page-content',
        '.text-content'
    ]

    for selector in content_selectors:
        content_elements = response.css(selector)
        if content_elements:
            content = content_elements.css('::text').getall()
            if content:
                combined_content = ' '.join(content)
                if len(combined_content.strip()) > 100:  # Minimum content length
                    return clean_content(combined_content)

    # Fallback: extract all paragraph text
    paragraphs = response.css('p::text').getall()
    if paragraphs:
        content = ' '.join(paragraphs)
        if len(content.strip()) > 100:
            return clean_content(content)

    # Final fallback: extract all text from body
    body_text = response.css('body ::text').getall()
    combined_text = ' '.join(body_text)
    return clean_content(combined_text)

def clean_content(content):
    """Clean and process scraped content."""
    if not content:
        return ''

    # Remove excessive whitespace
    content = re.sub(r'\s+', ' ', content)

    # Remove navigation and footer content
    content = re.sub(r'\b(home|menu|navigation|footer|copyright|privacy|terms|contact|about|login|signup|register|search)\b',
                    '', content, flags=re.IGNORECASE)

    # Remove emails, URLs, and phone numbers
    content = re.sub(r'\S+@\S+\.\S+', '[EMAIL]', content)
    content = re.sub(r'https?://[^\s]+', '[URL]', content)
    content = re.sub(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', '[PHONE]', content)

    # Remove non-printable characters
    content = re.sub(r'[^\x20-\x7E\n\r\t]', '', content)

    return content.strip()


def scrape_website_with_requests(url, max_pages=10):
    """Scrape content from a website and its subpaths using the new requests-based approach."""
    try:
        print(f"Starting requests-based crawl for: {url} (max {max_pages} pages)")

        # Parse the URL to get domain
        parsed_url = urlparse(url)
        domain = parsed_url.netloc
        base_url = f"{parsed_url.scheme}://{domain}"

        print(f"Crawling domain: {domain}")

        visited_urls = set()
        pages_to_visit = [url]
        all_pages = []
        pages_crawled = 0
        start_time = time.time()

        headers = {
            'User-Agent': 'DocifyBot/1.0 (https://docify.app)',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'Accept-Encoding': 'gzip, deflate, br',
        }

        while pages_to_visit and pages_crawled < max_pages and time.time() - start_time < 300:  # 5 minute timeout
            current_url = pages_to_visit.pop(0)

            if current_url in visited_urls:
                continue

            visited_urls.add(current_url)
            pages_crawled += 1

            print(f"Crawling page {pages_crawled}/{max_pages}: {current_url}")

            try:
                # Use smart fetch approach
                html_content, fetch_method, tried_rendering = smart_fetch_html(current_url, timeout=15)

                if not html_content:
                    print(f"Could not fetch content from: {current_url}")
                    continue

                # Detect content type
                response = requests.head(current_url, headers=headers, timeout=10)
                content_type = detect_content_type_from_response(response, current_url)

                # Create a mock response object for extract_content_by_type
                class MockResponse:
                    def __init__(self, content, headers, url):
                        self.content = content
                        self.headers = headers
                        self.url = url

                mock_response = MockResponse(html_content.encode('utf-8'), response.headers, current_url)

                # Extract content
                page_data = extract_content_by_type(mock_response, content_type, current_url)

                if not page_data:
                    print(f"Could not extract content from: {current_url}")
                    continue

                # Skip pages with very little content
                content_text = page_data.get('content', '')
                if not content_text or len(content_text.strip()) < 50:
                    print(f"Skipping page with insufficient content: {current_url}")
                    continue

                # Store the page data
                all_pages.append(page_data)

                # Find links to follow (only if we haven't reached max pages and content is HTML)
                if pages_crawled < max_pages and content_type == 'html':
                    try:
                        soup = BeautifulSoup(html_content, 'lxml')
                        for link in soup.find_all('a', href=True):
                            next_url = urljoin(current_url, link['href'])

                            # Only follow links within the same domain
                            if urlparse(next_url).netloc == domain:
                                # Skip common non-content URLs
                                path = urlparse(next_url).path.lower()
                                if not any(skip in path for skip in ['/search', '/login', '/admin', '/wp-admin', '/api/', '/feed', '/tag/', '/category/', '/author/']):
                                    if next_url not in visited_urls and next_url not in pages_to_visit:
                                        pages_to_visit.append(next_url)

                                        # Limit the queue size to prevent memory issues
                                        if len(pages_to_visit) > max_pages * 3:
                                            break
                    except Exception as link_error:
                        print(f"Error finding links on {current_url}: {link_error}")

                # Small delay to be respectful
                time.sleep(1.0)

            except Exception as page_error:
                print(f"Error crawling {current_url}: {page_error}")
                continue

        print(f"Requests-based crawling completed: {len(all_pages)} pages with content")

        if not all_pages:
            print("No pages were successfully crawled, falling back to single page")
            return scrape_single_page_fallback(url)

        # Limit to max_pages even if more were crawled
        all_pages = all_pages[:max_pages]

        # Analyze content types
        content_type_stats = {}
        for page in all_pages:
            content_type = page.get('content_type', 'unknown')
            content_type_stats[content_type] = content_type_stats.get(content_type, 0) + 1

        print(f"Content types found: {content_type_stats}")

        # Aggregate content from all pages
        total_word_count = sum(page['word_count'] for page in all_pages if page.get('word_count'))

        # Create combined content (limit content length to prevent memory issues)
        combined_content = []
        max_content_length = 75000
        current_length = 0

        # Group pages by content type for better organization
        pages_by_type = {}
        for page in all_pages:
            content_type = page.get('content_type', 'unknown')
            if content_type not in pages_by_type:
                pages_by_type[content_type] = []
            pages_by_type[content_type].append(page)

        # Add content by type
        for content_type, pages in pages_by_type.items():
            if current_length >= max_content_length:
                break

            # Add type header
            type_header = f"\n{'='*50}\n{content_type.upper()} CONTENT ({len(pages)} files)\n{'='*50}\n"
            if current_length + len(type_header) <= max_content_length:
                combined_content.append(type_header)
                current_length += len(type_header)

            for page in pages:
                if page.get('content') and current_length < max_content_length:
                    # Include metadata for non-HTML content
                    metadata_info = ""
                    if page.get('metadata'):
                        metadata_info = f"\n[Metadata: {page['metadata']}]"

                    page_content = f"\n--- {page['title']} ({page['url']}){metadata_info} ---\n{page['content']}"
                    if current_length + len(page_content) <= max_content_length:
                        combined_content.append(page_content)
                        current_length += len(page_content)
                    else:
                        # Truncate if it would exceed limit
                        remaining_space = max_content_length - current_length
                        if remaining_space > 100:
                            truncated_content = page_content[:remaining_space-10] + "..."
                            combined_content.append(truncated_content)
                        break

        # Get the main page data (first page crawled)
        main_page = all_pages[0] if all_pages else {}

        result = {
            'title': main_page.get('title', 'Multi-Type Content'),
            'description': f'Crawled {len(all_pages)} items from {domain} - Types: {", ".join([f"{k}({v})" for k,v in content_type_stats.items()])}',
            'content': '\n'.join(combined_content),
            'url': url,
            'word_count': total_word_count,
            'pages_crawled': len(all_pages),
            'scraped_at': 'now',
            'content_types': content_type_stats,
            'subpages': [{'url': p['url'], 'title': p['title'], 'word_count': p['word_count'], 'content_type': p.get('content_type', 'unknown')} for p in all_pages[1:]]
        }

        print(f"Successfully scraped {len(all_pages)} items with {total_word_count} total words")
        print(f"Content breakdown: {content_type_stats}")
        return result

    except Exception as error:
        print(f'Requests-based crawling error: {error}')
        print("Falling back to single page scraping...")
        return scrape_single_page_fallback(url)

def extract_excel_content(response, url):
    """Extract content from Excel files."""
    try:
        excel_file = BytesIO(response.content)
        df = pd.read_excel(excel_file)

        # Convert DataFrame to readable text
        content = []
        content.append(f"Sheet: {df.columns.tolist()}")
        content.append("Data Preview:")
        content.append(str(df.head(20)))  # First 20 rows

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'Excel Spreadsheet - {len(df)} rows, {len(df.columns)} columns',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'excel',
            'metadata': {
                'rows': len(df),
                'columns': len(df.columns),
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting Excel content: {e}")
        return None

def extract_csv_content(response, url):
    """Extract content from CSV files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        csv_text = response.content.decode(encoding, errors='ignore')
        lines = csv_text.split('\n')[:50]  # First 50 lines

        content = []
        content.append("CSV Data Preview:")
        content.extend(lines[:20])  # First 20 lines

        full_content = '\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'CSV File - {len(lines)} lines',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'csv',
            'metadata': {
                'lines': len(lines),
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting CSV content: {e}")
        return None

def extract_json_content(response, url):
    """Extract content from JSON files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        json_text = response.content.decode(encoding, errors='ignore')
        json_data = json.loads(json_text)

        # Convert JSON to readable text
        content = []
        content.append("JSON Structure:")
        content.append(json.dumps(json_data, indent=2)[:2000])  # First 2000 chars

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'JSON Data - {len(json_text)} characters',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'json',
            'metadata': {
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting JSON content: {e}")
        return None

def extract_xml_content(response, url):
    """Extract content from XML/RSS files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        xml_text = response.content.decode(encoding, errors='ignore')

        # Check if it's an RSS/Atom feed
        if '<rss' in xml_text.lower() or '<feed' in xml_text.lower():
            return extract_feed_content(xml_text, url)

        # Regular XML parsing
        from xml.etree import ElementTree as ET
        root = ET.fromstring(xml_text)

        content = []
        content.append("XML Structure:")
        content.append(xml_to_text(root, level=0)[:2000])

        full_content = '\n\n'.join(content)

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'XML Document - {len(xml_text)} characters',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'xml',
            'metadata': {
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting XML content: {e}")
        return None

def extract_feed_content(xml_text, url):
    """Extract content from RSS/Atom feeds."""
    try:
        feed = feedparser.parse(xml_text)

        content = []
        content.append(f"Feed Title: {feed.feed.get('title', 'Unknown')}")
        content.append(f"Feed Description: {feed.feed.get('description', '')}")
        content.append(f"Total Entries: {len(feed.entries)}")

        # Add recent entries
        for i, entry in enumerate(feed.entries[:10]):
            content.append(f"\n--- Entry {i+1} ---")
            content.append(f"Title: {entry.get('title', '')}")
            content.append(f"Link: {entry.get('link', '')}")
            if 'summary' in entry:
                content.append(f"Summary: {entry.summary[:500]}...")

        full_content = '\n'.join(content)

        return {
            'url': url,
            'title': feed.feed.get('title', extract_title_from_url(url)),
            'description': f'RSS/Atom Feed - {len(feed.entries)} entries',
            'content': full_content,
            'word_count': len(full_content.split()),
            'content_type': 'feed',
            'metadata': {
                'entries': len(feed.entries),
                'feed_type': 'rss' if '<rss' in xml_text.lower() else 'atom'
            }
        }
    except Exception as e:
        print(f"Error extracting feed content: {e}")
        return None

def extract_text_content(response, url):
    """Extract content from plain text files."""
    try:
        # Try to detect encoding
        detected = chardet.detect(response.content)
        encoding = detected.get('encoding', 'utf-8')

        text_content = response.content.decode(encoding, errors='ignore')

        return {
            'url': url,
            'title': extract_title_from_url(url),
            'description': f'Text File - {len(text_content)} characters',
            'content': text_content[:10000],  # Limit to 10K chars
            'word_count': len(text_content.split()),
            'content_type': 'text',
            'metadata': {
                'encoding': encoding,
                'file_size': len(response.content)
            }
        }
    except Exception as e:
        print(f"Error extracting text content: {e}")
        return None

def extract_html_content(response, url):
    """Extract content from HTML pages."""
    try:
        html_content = response.content.decode('utf-8', errors='ignore')
        soup = BeautifulSoup(html_content, 'lxml')

        title = extract_title(soup, url)
        description = extract_description(soup)
        content = extract_main_content(soup)

        return {
            'url': url,
                        'title': title,
            'description': description,
            'content': content,
            'word_count': len(content.split()) if content else 0,
                        'content_type': 'html'
                    }
    except Exception as e:
        print(f"Error extracting HTML content: {e}")
        return None

def extract_title_from_url(url):
    """Extract a title from URL when no other title is available."""
    path = urlparse(url).path
    filename = path.split('/')[-1]
    if '.' in filename:
        filename = filename.split('.')[0]
    return filename.replace('-', ' ').replace('_', ' ').title() or 'Document'

def xml_to_text(element, level=0):
    """Convert XML element to readable text."""
    indent = "  " * level
    text = f"{indent}<{element.tag}>"

    if element.text and element.text.strip():
        text += f" {element.text.strip()}"

    for child in element:
        text += "\n" + xml_to_text(child, level + 1)

    if not element.text or not element.text.strip():
        text += f"{indent}</{element.tag}>"

    return text

def extract_title(soup, url):
    """Extract page title."""
    title_selectors = [
        'title',
        'h1',
        'h1 a',
        '[property="og:title"]',
        'meta[name="title"]'
    ]

    for selector in title_selectors:
        try:
            if selector == 'title':
                element = soup.title
                if element:
                    return element.string.strip()
            elif selector.startswith('h1'):
                element = soup.select_one(selector)
                if element:
                    return element.get_text().strip()
            else:
                element = soup.select_one(selector)
                if element:
                    if 'property' in selector or 'name' in selector:
                        return element.get('content', '').strip()
                    else:
                        return element.get_text().strip()
        except:
                continue

    return extract_title_from_url(url)

def extract_description(soup):
    """Extract page description."""
    desc_selectors = [
        'meta[name="description"]',
        'meta[property="og:description"]',
        'meta[name="twitter:description"]'
    ]

    for selector in desc_selectors:
        try:
            element = soup.select_one(selector)
            if element:
                return element.get('content', '').strip()
        except:
            continue

    return ''

def extract_main_content(soup):
    """Extract main content from the page."""
    # Try various content selectors in order of preference
    content_selectors = [
        'main',
        '[role="main"]',
        '.content',
        '.main-content',
        '#content',
        '#main',
        'article',
        '.article-content',
        '.post-content',
        '.entry-content',
        '.page-content',
        '.text-content'
    ]

    for selector in content_selectors:
        try:
            content_elements = soup.select(selector)
            if content_elements:
                content = []
                for element in content_elements:
                    text = element.get_text()
                    if text.strip():
                        content.append(text)

                if content:
                    combined_content = ' '.join(content)
                    if len(combined_content.strip()) > 100:  # Minimum content length
                        return clean_content_fallback(combined_content)
        except:
            continue

    # Fallback: extract all paragraph text
    try:
        paragraphs = soup.select('p')
        if paragraphs:
            content = []
            for p in paragraphs:
                text = p.get_text().strip()
                if text:
                    content.append(text)

            if content:
                combined_content = ' '.join(content)
                if len(combined_content.strip()) > 100:
                    return clean_content_fallback(combined_content)
    except:
        pass

    # Final fallback: extract all text from body
    try:
        body = soup.body
        if body:
            body_text = body.get_text()
            return clean_content_fallback(body_text)
    except:
        pass

    return ''

def scrape_single_page_fallback(url):
    """Fallback single page scraper using the new smart fetch approach."""
    try:
        print(f"Fallback scraping single page: {url}")

        # Use smart fetch approach
        html_content, fetch_method, tried_rendering = smart_fetch_html(url, timeout=30)

        if not html_content:
            raise Exception("Could not fetch content from URL")

        # Parse with BeautifulSoup
        soup = BeautifulSoup(html_content, 'lxml')

        title = soup.title.string if soup.title else extract_title_from_url(url)

        meta_desc = soup.find('meta', attrs={'name': 'description'})
        description = meta_desc.get('content', '') if meta_desc else ''

        # Try to find main content
        content_selectors = [
            'main', '[role="main"]', '.content', '.main-content',
            '#content', '#main', 'article', '.article-content',
            '.post-content', '.entry-content'
        ]

        main_content = ''
        for selector in content_selectors:
            try:
                element = soup.select_one(selector)
                if element and len(element.get_text().strip()) > 100:
                    main_content = element.get_text().strip()
                    break
            except:
                continue

        if not main_content:
            main_content = soup.body.get_text().strip() if soup.body else ''

        cleaned_content = clean_content_fallback(main_content)

        return {
            'title': title,
            'description': description,
            'content': cleaned_content,
            'url': url,
            'word_count': len(cleaned_content.split()),
            'scraped_at': 'now'
        }

    except Exception as error:
        print(f'Fallback scraping error: {error}')
        raise Exception(f'Failed to scrape website: {str(error)}')


def clean_content_fallback(content):
    """Clean content for fallback scraper."""
    if not content:
        return ''

    content = re.sub(r'\s+', ' ', content)
    content = re.sub(r'\b(home|menu|navigation|footer|copyright|privacy|terms)\b', '', content, flags=re.IGNORECASE)
    content = re.sub(r'\S+@\S+\.\S+', '[EMAIL]', content)
    content = re.sub(r'https?://[^\s]+', '[URL]', content)
    content = re.sub(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', '[PHONE]', content)

    return content.strip()


# Appwrite configuration helper
def get_env_var(name, default=None):
    """Get environment variable with error handling."""
    value = os.environ.get(name, default)
    if value is None:
        raise ValueError(f"Required environment variable {name} is not set")
    return value

def init_appwrite_client(req):
    """Initialize Appwrite client using dynamic API key from headers."""
    try:
        client = Client()
        client.set_endpoint(get_env_var('APPWRITE_FUNCTION_API_ENDPOINT'))
        client.set_project(get_env_var('APPWRITE_FUNCTION_PROJECT_ID'))

        # Get dynamic API key from headers (x-appwrite-key)
        api_key = req.headers.get('x-appwrite-key')
        if not api_key:
            raise ValueError("Dynamic API key not found in x-appwrite-key header")

        client.set_key(api_key)
        print("Appwrite client initialized successfully with dynamic API key")
        return client
    except Exception as e:
        print(f"Error initializing Appwrite client: {e}")
        raise

# Database configuration
DATABASE_ID = os.environ.get('DATABASE_ID', 'docify_db')
DOCUMENTS_COLLECTION_ID = os.environ.get('DOCUMENTS_COLLECTION_ID', 'documents_table')


def scrape_website(url, max_pages=20):
    """Scrape content from a website URL and its subpaths using requests-based approach."""
    return scrape_website_with_requests(url, max_pages)


def update_document_status(databases, document_id, status, scraped_content=None):
    """Update document status in consolidated database."""
    try:
        update_data = {'status': status}

        # Include all scraped content data in the consolidated document
        if scraped_content:
            update_data['scraped_content'] = scraped_content['content']
            update_data['word_count'] = scraped_content['word_count']
            update_data['title'] = scraped_content['title']
            update_data['error_message'] = None  # Clear any previous errors

        databases.update_document(
            DATABASE_ID,
            DOCUMENTS_COLLECTION_ID,
            document_id,
            update_data
        )

        print(f'Document {document_id} status updated to {status}')
    except Exception as error:
        print(f'Failed to update document status: {error}')
        raise error


def trigger_llm_analysis(databases, document_id, scraped_data):
    """Trigger LLM analysis by updating document status (consolidated schema)."""
    try:
        # Just update the document status - the LLM analyzer will be triggered by event
        update_data = {
            'status': 'analyzing',
            'analysis_summary': 'Analysis in progress...',  # Placeholder
            'analysis_blocks': '[]'  # Empty JSON array as string
        }

        databases.update_document(
            DATABASE_ID,
            DOCUMENTS_COLLECTION_ID,
            document_id,
            update_data
        )

        print(f'LLM analysis triggered for document {document_id} (consolidated schema)')
    except Exception as error:
        print(f'Failed to trigger LLM analysis: {error}')
        raise error


def main(context):
    """Main function handler."""
    document_id = None
    url = None

    try:
        print('=== DOCUMENT SCRAPER STARTED ===')

        # Check if this is triggered by document creation (not analysis results)
        trigger_type = context.req.headers.get('x-appwrite-trigger', 'unknown')
        print(f'Trigger type: {trigger_type}')

        # Debug: Log all available headers
        print(f'Available headers: {list(context.req.headers.keys())}')
        for header_name, header_value in context.req.headers.items():
            print(f'Header {header_name}: {header_value}')

        if trigger_type == 'event':
            # Check if this is a document creation event
            collection_id = context.req.headers.get('x-appwrite-collection', '')
            event_type = context.req.headers.get('x-appwrite-event', '')
            print(f'Event triggered by collection: {collection_id}')
            print(f'Event type: {event_type}')

            # Extract collection name from event string if collection header is empty
            if not collection_id and 'collections.' in event_type:
                # Parse collection from event: databases.db.collections.collection_name.documents.doc_id.create
                try:
                    parts = event_type.split('.')
                    if len(parts) >= 4 and parts[2] == 'collections':
                        collection_id = parts[3]  # collection_name
                        print(f'Extracted collection from event: {collection_id}')
                except Exception as e:
                    print(f'Could not extract collection from event: {e}')

            # Only process document CREATION events from documents_table
            if collection_id != DOCUMENTS_COLLECTION_ID:
                print(f'Skipping: Event from {collection_id}, not from documents collection')
                return context.res.json({
                    'success': True,
                    'message': 'Skipped: Not a document event'
                }, 200)

            # With tables.*.rows.*.create trigger, check if this is actually a CREATE operation
            # We need to check the event payload to determine if it's a create or update
            is_create_operation = False

            # Check if this is a new document creation by looking for required fields
            if hasattr(context.req, 'body') and context.req.body:
                if isinstance(context.req.body, dict):
                    # For creation, we expect url and instructions to be present
                    if context.req.body.get('url') and context.req.body.get('instructions'):
                        is_create_operation = True
                        print('Detected CREATE operation based on request body')
                    else:
                        print('Request body missing url/instructions - not a creation event')

                elif isinstance(context.req.body, str):
                    import json
                    try:
                        body_data = json.loads(context.req.body)
                        if body_data.get('url') and body_data.get('instructions'):
                            is_create_operation = True
                            print('Detected CREATE operation from JSON body')
                        else:
                            print('JSON body missing url/instructions - not a creation event')
                    except Exception as e:
                        print(f'Could not parse JSON body: {e}')

            if not is_create_operation:
                print('Skipping: Not a document creation operation')
                return context.res.json({
                    'success': True,
                    'message': 'Skipped: Not a document creation event'
                }, 200)

        # Initialize Appwrite client with dynamic API key
        client = init_appwrite_client(context.req)
        databases = Databases(client)

        # Extract data based on trigger type
        if trigger_type == 'event':
            print('Processing event trigger')
            if context.req.body and context.req.body.get('$id') and context.req.body.get('url'):
                document_id = context.req.body['$id']
                url = context.req.body['url']
                print(f'Event data extracted - ID: {document_id}, URL: {url}')
            else:
                raise Exception('Event data missing required fields ($id or url)')
        elif trigger_type == 'http':
            print('Processing HTTP trigger')
            if context.req.body and context.req.body.get('documentId') and context.req.body.get('url'):
                document_id = context.req.body['documentId']
                url = context.req.body['url']
                print(f'API data extracted - ID: {document_id}, URL: {url}')
            else:
                raise Exception('API request missing required fields (documentId or url)')
        else:
            print('Processing unknown trigger type, attempting both extraction methods')
            if context.req.body:
                if context.req.body.get('$id') and context.req.body.get('url'):
                    document_id = context.req.body['$id']
                    url = context.req.body['url']
                    print(f'Extracted using event format - ID: {document_id}, URL: {url}')
                elif context.req.body.get('documentId') and context.req.body.get('url'):
                    document_id = context.req.body['documentId']
                    url = context.req.body['url']
                    print(f'Extracted using API format - ID: {document_id}, URL: {url}')
                else:
                    print(f'Request body structure: {context.req.body}')
                    raise Exception('Unable to extract documentId and url from request')
            else:
                raise Exception('No request body found')

        # Validate extracted data
        if not document_id or not url:
            raise Exception(f'Missing required data - documentId: {bool(document_id)}, url: {bool(url)}')

        print(f'Final data - Document ID: {document_id}, URL: {url}')

        # Validate URL format and check if domain exists
        print(f'Validating URL: {url}')
        try:
            parsed = urlparse(url)
            if not parsed.scheme or not parsed.netloc:
                raise Exception('Invalid URL format')

            # Basic domain validation
            domain = parsed.netloc.lower()
            if not domain or len(domain.split('.')) < 2:
                raise Exception('Invalid domain format')

            # Quick connectivity check
            import socket
            try:
                socket.gethostbyname(domain)
                print(f'Domain {domain} resolves successfully')
            except socket.gaierror as dns_error:
                print(f'Warning: Domain {domain} does not resolve: {dns_error}')
                # Don't fail completely, but log the warning

            print('URL validation passed')
        except Exception as url_error:
            print(f'URL validation failed: {url_error}')
            raise Exception(f'Invalid URL format: {str(url_error)}')

        # Update document status to scraping
        print('Updating document status to scraping...')
        update_document_status(databases, document_id, 'scraping')

        # Scrape the website with subpaths
        max_pages = int(os.environ.get('MAX_PAGES_TO_CRAWL', '20'))
        print(f'Starting website scraping for: {url} (max {max_pages} pages)')
        scraped_data = scrape_website(url, max_pages)

        print(f'Scraping completed - {scraped_data["word_count"]} words scraped from {scraped_data.get("pages_crawled", 1)} pages')

        # Update document with scraped content
        print('Updating document with scraped content...')
        update_document_status(databases, document_id, 'analyzing', scraped_data)

        # Trigger LLM analysis
        print('Triggering LLM analysis...')
        trigger_llm_analysis(databases, document_id, scraped_data)

        print('=== DOCUMENT SCRAPER COMPLETED SUCCESSFULLY ===')

        return context.res.json({
            'success': True,
            'message': 'Document scraped and analysis triggered successfully',
            'data': {
                'documentId': document_id,
                'wordCount': scraped_data['word_count'],
                'title': scraped_data['title'],
                'url': url
            }
        }, 200)

    except Exception as err:
        print(f'Document scraper failed: {err}')

        # Try to update document status to failed if we have a documentId
        if document_id:
            try:
                print(f'Updating document {document_id} status to failed')
                update_document_status(databases, document_id, 'failed')
            except Exception as update_error:
                print(f'Failed to update document status: {update_error}')

        print('=== DOCUMENT SCRAPER FAILED ===')

        return context.res.json({
            'success': False,
            'error': str(err),
            'documentId': document_id,
            'url': url
        }, 500)
